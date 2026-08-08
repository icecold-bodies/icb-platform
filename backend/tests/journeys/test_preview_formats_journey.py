"""v1.44 §3.4 — preview/export formats + Nadie's price permission (R1–R6).

Journeys (marker J144PF; purge at setup AND teardown):

1. R1 — a role-'full' user (Nadie's role) edits a BOM price from the costings
   page embed: right-click the row → "Edit permanently (this section)" → save →
   NO "Admin access required"; the value sticks and a BomOverrideHistory audit
   row exists.
2. R3/R4 — the embed header's Preview button opens the options dialog; pick
   Word + two ratios → a .docx downloads whose text carries the
   "Testing — {body} ({len} m)" heading and one "TOTAL COST @ {r}" line per
   ratio (never a combined total).
3. R5b — the approved /results page's ⬇ Export opens the options dialog; pick
   PDF + "line items" → a .pdf downloads (%PDF magic).
4. R5b/R6 — the costings dashboard row and CostingDetail show
   "{body type} ({length} m)"; the detail's Export button opens the shared
   dialog and confirm downloads through the approved endpoint.
5. R6 — the calculator's "Now costing body type" banner carries the entered
   length live: changing #f-length re-renders the suffix immediately.

Selector policy: Jinja pages by element ID; React by data-testid. No
wait_for_function (CSP has no unsafe-eval).
"""
from __future__ import annotations

import datetime as _dt
import io
import json

import pytest
from playwright.sync_api import Page, expect

from _common import admin_session, role_session, shot  # noqa: E402  (sys.path set in conftest)

T = 15_000
JOURNEY = "preview_formats"
MARK = "J144PF"
FULL_USER = "journey_full"


def _purge(db) -> None:
    from sqlalchemy import text
    db.execute(text(
        "DELETE FROM icb_costings.bom_override_history WHERE bom_id IN "
        "(SELECT id FROM icb_costings.bill_of_materials WHERE material_id IN "
        " (SELECT id FROM icb_costings.materials WHERE name LIKE :m))"), {"m": f"{MARK}%"})
    db.execute(text(
        "DELETE FROM icb_costings.calculations WHERE trailer_type_id IN "
        "(SELECT id FROM icb_costings.trailer_types WHERE name LIKE :m)"), {"m": f"{MARK}%"})
    db.execute(text(
        "DELETE FROM icb_costings.bill_of_materials WHERE material_id IN "
        "(SELECT id FROM icb_costings.materials WHERE name LIKE :m)"), {"m": f"{MARK}%"})
    db.execute(text("DELETE FROM icb_costings.materials WHERE name LIKE :m"), {"m": f"{MARK}%"})
    db.execute(text("DELETE FROM icb_costings.trailer_types WHERE name LIKE :m"), {"m": f"{MARK}%"})
    db.commit()


@pytest.fixture(scope="module")
def staged():
    """Calculable marker trailer (default_length 6.0) + one plain BOM row +
    a pending costing record + the role-'full' journey user."""
    from app.database import (BillOfMaterial, CalculationRecord, Material,
                              SessionLocal, TrailerType, User)
    created_user = False
    with SessionLocal() as db:
        _purge(db)
        trailer = TrailerType(name=f"{MARK} FORMATS BODY", is_active=True,
                              default_length=6.0, default_width=2.4, default_height=2.4)
        db.add(trailer)
        db.flush()
        mat = Material(name=f"{MARK} PANEL SHEET", unit_of_measure="m2", price_per_unit=250.0)
        db.add(mat)
        db.flush()
        bom = BillOfMaterial(trailer_type_id=trailer.id, material_id=mat.id,
                             formula_expression="1", waste_percentage=0,
                             bom_section="PANELS", sort_order=1)
        db.add(bom)
        db.flush()
        admin = db.query(User).filter_by(username="admin").first()
        rec = CalculationRecord(
            trailer_type_id=trailer.id, user_id=admin.id, status="pending",
            dimensions_json=json.dumps({"length": 6.0, "width": 2.4, "height": 2.4}),
            result_json=json.dumps({
                "items": [{"category": "PANELS", "material": f"{MARK} PANEL SHEET",
                           "material_code": "", "formula": "1", "quantity": 1.0,
                           "unit": "m2", "unit_price": 250.0, "waste_pct": 0,
                           "line_cost": 250.0, "last_updated": None}],
                "category_totals": {"PANELS": 250.0},
                "grand_total": 250.0, "cost_per_sqm": 17.36,
                # The /results Jinja page iterates result.geometry unguarded —
                # a saved record without it 500s (banked: enriched-seed trap).
                "geometry": {"floor_area": 14.4, "surface_area": 71.0,
                             "wall_area": 28.8, "roof_area": 14.4},
            }))
        db.add(rec)
        if db.query(User).filter_by(username=FULL_USER).first() is None:
            db.add(User(username=FULL_USER, password_hash="x", role="full"))
            created_user = True
        db.commit()
        ids = {"trailer": trailer.id, "rec": rec.id, "bom": bom.id, "name": trailer.name}
    yield ids
    with SessionLocal() as db:
        _purge(db)
        if created_user:
            u = db.query(User).filter_by(username=FULL_USER).first()
            if u is not None:
                db.delete(u)
        db.commit()


def _record_count() -> int:
    from app.database import CalculationRecord, SessionLocal
    with SessionLocal() as db:
        return db.query(CalculationRecord).count()


def _docx_text(docx_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "|".join(parts)


# ── 1. R1: full-role permanent price save with audit ──────────────────────────
def test_full_role_price_save_from_costings_page(page: Page, live_server: str, staged) -> None:
    ids = staged
    role_session(page, FULL_USER, base=live_server)
    page.goto("/mes-app/costings/new")

    frame = page.frame_locator("iframe[title='Calculator (live costing app)']")
    expect(frame.locator("#trailer-select")).to_be_visible(timeout=30_000)
    frame.locator("#trailer-select").select_option(str(ids["trailer"]))

    row = frame.locator("[data-material-id]").first
    expect(row).to_be_visible(timeout=30_000)
    row.click(button="right")
    frame.locator(".ctx-menu-item", has_text="Edit permanently (this section)").click()

    modal = frame.locator("#modal-section-price")
    expect(modal).to_be_visible(timeout=T)
    frame.locator("#sprice-input").fill("199.99")
    shot(page, "01-full-role-price-modal", journey=JOURNEY)
    frame.locator("#sprice-save-btn").click()

    # The save toasts success — NOT "Save failed: Admin access required".
    expect(frame.locator("#toast-container .toast-msg",
                         has_text="Section price saved").first).to_be_visible(timeout=T)
    expect(frame.locator("#toast-container .toast-msg",
                         has_text="Admin access required")).to_have_count(0)
    shot(page, "02-full-role-price-saved", journey=JOURNEY)

    from app.database import BillOfMaterial, BomOverrideHistory, SessionLocal
    with SessionLocal() as db:
        bom = db.query(BillOfMaterial).filter_by(id=ids["bom"]).first()
        assert bom.unit_price_override == 199.99, "override did not stick"
        audit = (db.query(BomOverrideHistory).filter_by(bom_id=ids["bom"])
                 .order_by(BomOverrideHistory.id.desc()).first())
        assert audit is not None and audit.new_price == 199.99, "audit row missing"


# ── 2. R3/R4: preview dialog → Word with two ratios ───────────────────────────
def test_preview_dialog_word_two_ratios(page: Page, live_server: str, staged, tmp_path) -> None:
    ids = staged
    admin_session(page, base=live_server)
    page.goto("/mes-app/costings/new")

    frame = page.frame_locator("iframe[title='Calculator (live costing app)']")
    expect(frame.locator("#trailer-select")).to_be_visible(timeout=30_000)
    frame.locator("#trailer-select").select_option(str(ids["trailer"]))
    frame.locator("#f-length").fill("6.5")
    expect(frame.locator("#approve-btn")).to_be_enabled(timeout=30_000)

    before = _record_count()
    page.get_by_test_id("preview-btn").click()
    expect(page.get_by_test_id("export-confirm")).to_be_visible(timeout=T)
    page.get_by_test_id("export-fmt-word").click()
    page.get_by_test_id("export-ratio-35%").check()
    page.get_by_test_id("export-ratio-55%").check()
    shot(page, "03-word-two-ratios-dialog", journey=JOURNEY)

    with page.expect_download(timeout=T) as dl_info:
        page.get_by_test_id("export-confirm").click()
    download = dl_info.value
    today = _dt.date.today().strftime("%Y-%m-%d")
    assert download.suggested_filename == f"Testing - {ids['name']} - {today}.docx", \
        download.suggested_filename

    target = tmp_path / download.suggested_filename
    download.save_as(target)
    text = _docx_text(target.read_bytes())
    assert f"Testing — {ids['name']} (6.5 m)" in text
    assert "TOTAL COST @ 35%" in text and "TOTAL COST @ 55%" in text
    assert text.count("TOTAL COST @") == 2            # one line per ratio, never combined
    assert "Client:  — no client selected —" in text
    assert f"{MARK} PANEL SHEET" not in text          # default detail = totals only
    assert _record_count() == before, "preview must write NOTHING to the DB"
    shot(page, "04-word-downloaded", journey=JOURNEY)


# ── 3. R5b: approved /results page → Export dialog → PDF with line items ──────
def test_results_export_dialog_pdf_items(page: Page, live_server: str, staged, tmp_path) -> None:
    ids = staged
    admin_session(page, base=live_server)
    page.goto(f"/results/{ids['rec']}")

    page.locator("#excel-export-btn").click()
    expect(page.locator("#modal-export-options")).to_be_visible(timeout=T)
    page.locator("#exp-fmt-pdf").click()
    page.locator("input[name='exp-detail'][value='items']").check()
    shot(page, "05-results-export-dialog", journey=JOURNEY)

    with page.expect_download(timeout=T) as dl_info:
        page.locator("#exp-confirm-btn").click()
    download = dl_info.value
    assert download.suggested_filename.endswith(".pdf"), download.suggested_filename

    target = tmp_path / download.suggested_filename
    download.save_as(target)
    assert target.read_bytes()[:5] == b"%PDF-"
    shot(page, "06-results-pdf-downloaded", journey=JOURNEY)


# ── 4. R5b/R6: dashboard + detail length display; detail Export downloads ─────
def test_costing_detail_export_and_length_display(page: Page, live_server: str, staged, tmp_path) -> None:
    ids = staged
    admin_session(page, base=live_server)
    page.goto("/mes-app/costings")

    # Dashboard row shows "{body type} ({length} m)" (R6).
    row_cell = page.get_by_text(f"{ids['name']} (6 m)").first
    expect(row_cell).to_be_visible(timeout=30_000)
    shot(page, "07-dashboard-length", journey=JOURNEY)

    row_cell.click()
    # Detail header + Configuration card carry the suffix too.
    expect(page.get_by_text(f"{ids['name']} (6 m)").first).to_be_visible(timeout=30_000)

    page.get_by_test_id("export-open-btn").click()
    expect(page.get_by_test_id("export-confirm")).to_be_visible(timeout=T)
    page.get_by_test_id("export-fmt-word").click()
    shot(page, "08-detail-export-dialog", journey=JOURNEY)
    with page.expect_download(timeout=T) as dl_info:
        page.get_by_test_id("export-confirm").click()
    download = dl_info.value
    assert download.suggested_filename.endswith(".docx"), download.suggested_filename
    target = tmp_path / download.suggested_filename
    download.save_as(target)
    text = _docx_text(target.read_bytes())
    assert f"#{ids['rec']} — {ids['name']} (6 m)" in text   # approved heading, no "Testing"
    assert "Testing" not in text
    shot(page, "09-detail-word-downloaded", journey=JOURNEY)


# ── 5. R6: banner carries the entered length live ─────────────────────────────
def test_banner_length_live(page: Page, live_server: str, staged) -> None:
    ids = staged
    admin_session(page, base=live_server)
    page.goto("/calculator")

    expect(page.locator("#trailer-select")).to_be_visible(timeout=T)
    page.locator("#trailer-select").select_option(str(ids["trailer"]))
    banner = page.locator("#topbar-title")
    expect(banner).to_contain_text(f"{ids['name']} (6 m)", timeout=30_000)
    shot(page, "10-banner-default-length", journey=JOURNEY)

    page.locator("#f-length").fill("7.5")
    expect(banner).to_contain_text(f"{ids['name']} (7.5 m)", timeout=T)
    shot(page, "11-banner-live-length", journey=JOURNEY)
