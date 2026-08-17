"""WO v1.47 lane B — the end-user snapshot on a costing + the document lines (0040).

Covers: the _end_user_snapshot resolver (fields, cross-customer 422, inactive 422, null
passthrough); the snapshot serialized on GET /api/calculations/{id} and the list; the
snapshot FROZEN against later edits to the end-user record; FK ondelete=SET NULL keeping
the display values through a hard delete; and the shared document context putting
"End user: …" / "End user contact: …" into ALL THREE formats — with the optional path
(no end user → no lines at all, in every format) asserted alongside.

The contact snapshot is deliberately re-asserted where the two sit side by side: this WO
extends it, and must never alter its values.

House pattern (test_calculation_contact_api.py): live test DB, marker rows 'V147EU*',
module-local fixtures, purge on both sides.
"""
import io

import pytest

_MARK = "V147EUS"

RESULT_FIXTURE = {
    "items": [
        {"category": "FLOOR", "material": "EU FLOOR SHEET", "material_code": "EF-1",
         "formula": "L*W", "quantity": 10.0, "unit": "m2", "unit_price": 100.0,
         "waste_pct": 0, "line_cost": 1000.0, "last_updated": None},
    ],
    "category_totals": {"FLOOR": 1000.0},
    "grand_total": 1000.0,
    "profit_margin": 10,
    "profit_amount": 100.0,
}
DIMS_FIXTURE = {"length": 7.5, "width": 2.4, "height": 2.4}


def _purge(db) -> None:
    from sqlalchemy import text
    db.execute(text(
        "DELETE FROM icb_costings.calculations cal USING icb_costings.customers c "
        "WHERE cal.customer_id = c.id AND c.name LIKE :m"), {"m": f"{_MARK}%"})
    db.execute(text(
        "DELETE FROM icb_costings.customer_end_users eu USING icb_costings.customers c "
        "WHERE eu.customer_id = c.id AND c.name LIKE :m"), {"m": f"{_MARK}%"})
    db.execute(text("DELETE FROM icb_costings.customers WHERE name LIKE :m"), {"m": f"{_MARK}%"})
    db.commit()


@pytest.fixture(scope="module")
def app_mod():
    import app.main as m
    from starlette.testclient import TestClient
    with TestClient(m.app):
        yield m


@pytest.fixture
def api(app_mod):
    """Admin client. The calculator endpoints authenticate via get_current_user(request, db)
    INSIDE the handler, which no dependency override reaches — so also mint a REAL
    UserSession row and send it as a raw Cookie header ([[testclient-session-cookie]]:
    raw header, not the httpx jar)."""
    import uuid
    from app.database import SessionLocal, User, UserSession
    from app.deps import require_admin, require_user
    from starlette.testclient import TestClient
    sid = str(uuid.uuid4())
    with SessionLocal() as db:
        _purge(db)
        admin = db.query(User).filter_by(username="admin").first()
        db.add(UserSession(id=sid, user_id=admin.id, role=admin.role, expires_at=None))
        db.commit()
    app_mod.app.dependency_overrides[require_user] = lambda: admin
    app_mod.app.dependency_overrides[require_admin] = lambda: admin
    with TestClient(app_mod.app) as c:
        c.headers["Cookie"] = f"session_id={sid}"
        yield c
    app_mod.app.dependency_overrides.pop(require_user, None)
    app_mod.app.dependency_overrides.pop(require_admin, None)
    with SessionLocal() as db:
        db.query(UserSession).filter_by(id=sid).delete()
        db.commit()
        _purge(db)


@pytest.fixture
def cust_with_end_users():
    """A reseller customer with two end users (one primary) + one soft-deleted, plus a
    second customer that owns none (the cross-customer 422 case)."""
    from app.database import Customer, CustomerEndUser, SessionLocal
    with SessionLocal() as db:
        _purge(db)
        reseller = Customer(name=f"{_MARK} Reseller Ltd", bp_code=f"{_MARK}1", is_active=True)
        direct = Customer(name=f"{_MARK} Direct Ltd", bp_code=f"{_MARK}2", is_active=True)
        db.add_all([reseller, direct])
        db.flush()
        acme = CustomerEndUser(customer_id=reseller.id, company_name=f"{_MARK} ACME Foods",
                               contact_name="Thabo Nkosi", contact_role="Fleet",
                               contact_email="thabo@acme.co", contact_telephone="011 999",
                               is_primary=True, active=True)
        bare = CustomerEndUser(customer_id=reseller.id, company_name=f"{_MARK} Bare Co",
                               active=True)
        dead = CustomerEndUser(customer_id=reseller.id, company_name=f"{_MARK} Gone Co",
                               active=False)
        db.add_all([acme, bare, dead])
        db.commit()
        ids = {"reseller": reseller.id, "direct": direct.id,
               "acme": acme.id, "bare": bare.id, "dead": dead.id}
    yield ids
    with SessionLocal() as db:
        _purge(db)


def _mk_calc(db, cust_id, end_user=None):
    """A minimal saved costing row (the /api/approve pipeline needs a full BOM universe the
    test DB doesn't have — the resolver is covered at unit level above, persistence and
    serialization at the record level here, and the whole pipeline in the journey)."""
    from app.database import CalculationRecord
    rec = CalculationRecord(
        customer_id=cust_id,
        dimensions_json="{}", result_json='{"items": []}', status="pending",
        end_user_id=end_user.id if end_user is not None else None,
        end_user_company=end_user.company_name if end_user is not None else None,
        end_user_contact_name=end_user.contact_name if end_user is not None else None,
        end_user_contact_email=end_user.contact_email if end_user is not None else None,
        end_user_contact_telephone=end_user.contact_telephone if end_user is not None else None,
        end_user_contact_role=end_user.contact_role if end_user is not None else None,
    )
    db.add(rec)
    db.commit()
    db.refresh(rec)
    return rec


# ── _end_user_snapshot resolver ───────────────────────────────────────────────

def test_end_user_snapshot_resolves_fields(api, cust_with_end_users):
    from app.database import SessionLocal
    from app.routers.calculator import _end_user_snapshot
    ids = cust_with_end_users
    with SessionLocal() as db:
        snap = _end_user_snapshot(db, ids["reseller"], ids["acme"])
    assert snap == {"end_user_id": ids["acme"], "end_user_company": f"{_MARK} ACME Foods",
                    "end_user_contact_name": "Thabo Nkosi",
                    "end_user_contact_email": "thabo@acme.co",
                    "end_user_contact_telephone": "011 999",
                    "end_user_contact_role": "Fleet"}


def test_end_user_snapshot_null_passthrough(api, cust_with_end_users):
    """The optional path: no end user picked → all-NULL columns, no error."""
    from app.database import SessionLocal
    from app.routers.calculator import _end_user_snapshot
    ids = cust_with_end_users
    empty = {"end_user_id": None, "end_user_company": None, "end_user_contact_name": None,
             "end_user_contact_email": None, "end_user_contact_telephone": None,
             "end_user_contact_role": None}
    with SessionLocal() as db:
        assert _end_user_snapshot(db, ids["reseller"], None) == empty
        assert _end_user_snapshot(db, None, ids["acme"]) == empty
        assert _end_user_snapshot(db, None, None) == empty


def test_end_user_snapshot_guards_422(api, cust_with_end_users):
    from fastapi import HTTPException
    from app.database import SessionLocal
    from app.routers.calculator import _end_user_snapshot
    ids = cust_with_end_users
    with SessionLocal() as db:
        with pytest.raises(HTTPException) as e1:                   # another customer's end user
            _end_user_snapshot(db, ids["direct"], ids["acme"])
        assert e1.value.status_code == 422
        with pytest.raises(HTTPException) as e2:                   # soft-deleted end user
            _end_user_snapshot(db, ids["reseller"], ids["dead"])
        assert e2.value.status_code == 422
        with pytest.raises(HTTPException) as e3:                   # nonexistent end user
            _end_user_snapshot(db, ids["reseller"], 99999999)
        assert e3.value.status_code == 422


def test_end_user_snapshot_leaves_the_contact_snapshot_alone(api, cust_with_end_users):
    """The two chokepoints are siblings, not one — neither writes the other's columns."""
    from app.database import SessionLocal
    from app.routers.calculator import _contact_snapshot, _end_user_snapshot
    ids = cust_with_end_users
    with SessionLocal() as db:
        eu = _end_user_snapshot(db, ids["reseller"], ids["acme"])
        ct = _contact_snapshot(db, ids["reseller"], None)
    assert not any(k.startswith("contact_") for k in eu)
    assert not any(k.startswith("end_user") for k in ct)
    assert all(v is None for v in ct.values())


# ── Persistence + serialization ───────────────────────────────────────────────

def test_snapshot_serialized_on_get_and_list(api, cust_with_end_users):
    from app.database import CustomerEndUser, SessionLocal
    ids = cust_with_end_users
    with SessionLocal() as db:
        acme = db.get(CustomerEndUser, ids["acme"])
        rec_id = _mk_calc(db, ids["reseller"], end_user=acme).id

    got = api.get(f"/api/calculations/{rec_id}").json()
    assert got["end_user_id"] == ids["acme"]
    assert got["end_user_company"] == f"{_MARK} ACME Foods"
    assert got["end_user_contact_name"] == "Thabo Nkosi"
    assert got["end_user_contact_email"] == "thabo@acme.co"
    assert got["end_user_contact_telephone"] == "011 999"
    assert got["end_user_contact_role"] == "Fleet"

    rows = api.get("/api/calculations", params={"limit": 50}).json()
    mine = next((r for r in rows if r["id"] == rec_id), None)
    assert mine is not None, "fresh record missing from the list"
    assert mine["end_user_company"] == f"{_MARK} ACME Foods"


def test_no_end_user_serializes_as_null(api, cust_with_end_users):
    from app.database import SessionLocal
    ids = cust_with_end_users
    with SessionLocal() as db:
        rec_id = _mk_calc(db, ids["reseller"], end_user=None).id
    got = api.get(f"/api/calculations/{rec_id}").json()
    assert got["end_user_id"] is None and got["end_user_company"] is None


def test_snapshot_frozen_against_later_edits(api, cust_with_end_users):
    """THE point of the snapshot: editing the end-user record afterwards must not rewrite
    what an already-saved costing says."""
    from app.database import CustomerEndUser, SessionLocal
    ids = cust_with_end_users
    with SessionLocal() as db:
        acme = db.get(CustomerEndUser, ids["acme"])
        rec_id = _mk_calc(db, ids["reseller"], end_user=acme).id
        acme.company_name = f"{_MARK} ACME RENAMED"
        acme.contact_name = "Someone Else"
        acme.contact_email = "renamed@acme.co"
        db.commit()

    got = api.get(f"/api/calculations/{rec_id}").json()
    assert got["end_user_company"] == f"{_MARK} ACME Foods"      # snapshot, not the live row
    assert got["end_user_contact_name"] == "Thabo Nkosi"
    assert got["end_user_contact_email"] == "thabo@acme.co"
    assert got["end_user_id"] == ids["acme"]                     # FK still points at the row


def test_end_user_hard_delete_sets_null_keeps_snapshot(api, cust_with_end_users):
    """ondelete=SET NULL — a hard DELETE nulls the FK but the snapshot columns keep the
    quote's history intact."""
    from sqlalchemy import text
    from app.database import CustomerEndUser, SessionLocal
    ids = cust_with_end_users
    with SessionLocal() as db:
        acme = db.get(CustomerEndUser, ids["acme"])
        rec_id = _mk_calc(db, ids["reseller"], end_user=acme).id
        db.execute(text("DELETE FROM icb_costings.customer_end_users WHERE id = :i"),
                   {"i": ids["acme"]})
        db.commit()

    got = api.get(f"/api/calculations/{rec_id}").json()
    assert got["end_user_id"] is None                            # FK went NULL, no cascade
    assert got["end_user_company"] == f"{_MARK} ACME Foods"      # history preserved
    assert got["end_user_contact_name"] == "Thabo Nkosi"


# ── Shared document context → all three formats ───────────────────────────────

def _ctx(company=None, contact=None):
    from app.services.document_context import build_doc_ctx
    return build_doc_ctx(
        mode="approved", heading="Q-1 — EU BODY (7.5 m)", sub="EU BODY",
        client_name="Reseller Ltd", spec_pairs=[("Length (m)", 7.5)], spec_options=[],
        result=dict(RESULT_FIXTURE), ratios=[0.55], detail="items", db=None,
        end_user_company=company, end_user_contact_name=contact)


def test_build_end_user_lines_is_the_single_wording_source():
    from app.services.document_context import build_end_user_lines
    assert build_end_user_lines("ACME Foods", "Thabo") == [
        "End user: ACME Foods", "End user contact: Thabo"]
    assert build_end_user_lines("ACME Foods", None) == ["End user: ACME Foods"]
    assert build_end_user_lines("ACME Foods", "   ") == ["End user: ACME Foods"]
    assert build_end_user_lines(None, None) == []
    assert build_end_user_lines("   ", "Thabo") == []      # a person with no company is not renderable


def _xlsx_flat(ctx) -> str:
    import openpyxl
    from app.routers.exports import _render_xlsx
    ws = openpyxl.load_workbook(io.BytesIO(_render_xlsx(ctx).getvalue())).active
    return "|".join(str(c.value) for row in ws.iter_rows() for c in row if c.value is not None)


def _docx_flat(ctx) -> str:
    from docx import Document
    from app.routers.exports import _render_docx
    doc = Document(io.BytesIO(_render_docx(ctx).getvalue()))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [cell.text for cell in row.cells]
    return "|".join(parts)


def _pdf_page1(ctx) -> str:
    from pypdf import PdfReader
    from app.routers.exports import _render_pdf
    return (PdfReader(io.BytesIO(_render_pdf(ctx))).pages[0].extract_text() or "").replace("\n", " ")


@pytest.mark.parametrize("flatten", [_xlsx_flat, _docx_flat, _pdf_page1],
                         ids=["excel", "word", "pdf"])
def test_every_format_prints_both_end_user_lines(flatten):
    text = flatten(_ctx("ACME Foods", "Thabo Nkosi"))
    assert "End user: ACME Foods" in text
    assert "End user contact: Thabo Nkosi" in text
    assert "Reseller Ltd" in text, "the client line must still be there"


@pytest.mark.parametrize("flatten", [_xlsx_flat, _docx_flat, _pdf_page1],
                         ids=["excel", "word", "pdf"])
def test_company_without_a_person_prints_one_line(flatten):
    text = flatten(_ctx("ACME Foods", None))
    assert "End user: ACME Foods" in text
    assert "End user contact" not in text


@pytest.mark.parametrize("flatten", [_xlsx_flat, _docx_flat, _pdf_page1],
                         ids=["excel", "word", "pdf"])
def test_no_end_user_prints_nothing_at_all(flatten):
    """The optional path across every format: not a blank line, not an empty label."""
    text = flatten(_ctx(None, None))
    assert "End user" not in text
    assert "Reseller Ltd" in text


def test_excel_layout_below_the_client_is_unmoved_without_an_end_user():
    """The xlsx renderer counts rows rather than hardcoding them, so the block can be
    absent — with no end user every row must land exactly where it did before this WO
    (client A3, section header A4, dimensions row 5)."""
    import openpyxl
    from app.routers.exports import _render_xlsx
    ws = openpyxl.load_workbook(io.BytesIO(_render_xlsx(_ctx(None, None)).getvalue())).active
    assert str(ws["A3"].value).startswith("Client:")
    assert ws["A4"].value == "DIMENSIONS & BODY OPTIONS"
    assert ws.cell(row=5, column=1).value == "Length (m)"

    # …and WITH an end user the same blocks simply shift down by the two lines.
    ws2 = openpyxl.load_workbook(
        io.BytesIO(_render_xlsx(_ctx("ACME Foods", "Thabo")).getvalue())).active
    assert str(ws2["A3"].value).startswith("Client:")
    assert ws2["A4"].value == "End user: ACME Foods"
    assert ws2["A5"].value == "End user contact: Thabo"
    assert ws2["A6"].value == "DIMENSIONS & BODY OPTIONS"
    assert ws2.cell(row=7, column=1).value == "Length (m)"


def test_approved_exports_read_the_snapshot_not_the_live_row(api, cust_with_end_users):
    """End to end through the real export route: change the end-user record after saving
    and the approved costing still prints what it was quoted with."""
    import openpyxl
    from app.database import CustomerEndUser, SessionLocal, TrailerType
    import json as _json
    ids = cust_with_end_users
    with SessionLocal() as db:
        acme = db.get(CustomerEndUser, ids["acme"])
        tt = db.query(TrailerType).first()
        rec = _mk_calc(db, ids["reseller"], end_user=acme)
        rec.trailer_type_id = tt.id if tt else None
        rec.dimensions_json = _json.dumps(DIMS_FIXTURE)
        rec.result_json = _json.dumps(RESULT_FIXTURE)
        db.commit()
        rec_id = rec.id
        acme.company_name = f"{_MARK} ACME RENAMED"
        db.commit()

    r = api.get(f"/results/{rec_id}/export/excel")
    assert r.status_code == 200, r.text
    ws = openpyxl.load_workbook(io.BytesIO(r.content)).active
    flat = "|".join(str(c.value) for row in ws.iter_rows() for c in row if c.value is not None)
    assert f"End user: {_MARK} ACME Foods" in flat
    assert "RENAMED" not in flat
