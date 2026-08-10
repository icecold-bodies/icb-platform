"""Costing document exports — Excel / Word / PDF, preview AND approved (v1.44).

Every document is drawn from ONE shared, layout-neutral context
(services/document_context.build_doc_ctx) so the three formats can't drift:
heading → client → dimensions & body options → CATEGORY TOTALS → PRICE SUMMARY
(one "TOTAL COST @ r" per selected ratio; never combining two ratios; no cost
per m²) → line items only when detail == "items".

Endpoints:
  POST /api/export/preview           — live result, format = excel|word|pdf
  POST /api/export/excel-preview     — v1.44 F1 alias (format forced to excel)
  GET  /results/{id}/export/excel    — approved (+ highlight= colour-coding)
  GET  /results/{id}/export/word     — approved (NEW)
  GET  /results/{id}/export/pdf      — approved
  GET  /results/{id}/report          — templated quote PDF (untouched surface)

All accept detail= ("totals"|"items", default "items" — the bare legacy URLs
keep their full line-item content) and ratios= (comma list of divisors; default
= the ratio saved/selected on the page). Previews still write NOTHING and
consume no quote number. The former WeasyPrint branch of the PDF path is gone:
it was dead-by-design in every target environment (requirements.txt / ADR 0017
deliberately omit weasyprint), and routing all environments through ReportLab
keeps layout parity a property of the code, not of installed native libs.
"""
import json
import re
from datetime import datetime, timezone, timedelta
from io import BytesIO

from fastapi import Request, APIRouter, Depends, HTTPException
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.orm import Session

from ..database import get_db, CalculationRecord, TrailerType, BillOfMaterial
from ..deps import get_current_user, user_can
from ..services import resolve_report_template, strip_excluded_items, _bom_load_options
from ..services.document_context import (
    VALID_DETAILS, VALID_FORMATS,
    body_type_with_length, build_doc_ctx, parse_ratios,
)

router = APIRouter()

_FORMAT_GATES = {"excel": "export.excel", "word": "export.word", "pdf": "export.pdf"}
_FORMAT_MEDIA = {
    "excel": ("xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    "word":  ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf":   ("pdf",  "application/pdf"),
}


def _require_export_user(request: Request, db: Session, fmt: str):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401)
    gate = _FORMAT_GATES[fmt]
    if not user_can(user, gate, db):
        raise HTTPException(status_code=403, detail=f"Permission denied: {gate}")
    return user


def _render_document(fmt: str, ctx: dict) -> BytesIO:
    if fmt == "excel":
        return _render_xlsx(ctx)
    if fmt == "word":
        return _render_docx(ctx)
    return BytesIO(_render_pdf(ctx))


# ── Excel renderer ────────────────────────────────────────────────────────────

def _render_xlsx(ctx: dict):
    """Render the shared doc ctx as a workbook (BytesIO). Extends the v1.44 F1
    builder to the ratified R2 order; the collapsible item outline, literal
    section-header totals and the highlight legend are retained."""
    import openpyxl
    import io
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    highlight = ctx["highlight"]
    override_materials = ctx["override_materials"]
    recently_updated_mats = ctx["recently_updated_mats"]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cost Breakdown"

    hdr_fill = PatternFill("solid", fgColor="1C2333")
    cat_fill = PatternFill("solid", fgColor="1F3A5F")
    total_fill = PatternFill("solid", fgColor="0D4A8A")
    grand_fill = PatternFill("solid", fgColor="388BFD")
    thin = Border(
        bottom=Side(style="thin", color="30363D"),
        right=Side(style="thin", color="30363D"),
    )

    # 1 — heading
    ws.merge_cells("A1:I1")
    t = ws["A1"]
    t.value = ctx["heading"]
    t.font = Font(bold=True, size=14, name="Calibri", color=ctx["heading_color"])
    t.fill = PatternFill("solid", fgColor="0D1117")
    t.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells("A2:I2")
    s = ws["A2"]
    s.value = ctx["sub"]
    s.font = Font(size=11, color="8B949E", name="Calibri")
    s.fill = PatternFill("solid", fgColor="161B22")
    s.alignment = Alignment(horizontal="center")

    # 2 — client (always present; previews without a customer show the placeholder)
    ws.merge_cells("A3:I3")
    c3 = ws["A3"]
    c3.value = f"Client:  {ctx['client_name']}"
    c3.font = Font(bold=True, size=11, color="0D1117", name="Calibri")
    c3.alignment = Alignment(horizontal="center")

    # 3 — dimensions + body options
    ws["A4"] = "DIMENSIONS & BODY OPTIONS"
    ws["A4"].font = Font(bold=True, color="388BFD", name="Calibri")
    for i, (lbl, val) in enumerate(ctx["spec_pairs"]):
        ws.cell(row=5, column=i * 2 + 1, value=lbl).font = Font(color="444444", name="Calibri")
        ws.cell(row=5, column=i * 2 + 2, value=val).font = Font(bold=True, color="000000", name="Calibri")
    row = 6
    for lbl, val in ctx["spec_options"]:
        ws.cell(row=row, column=1, value=lbl).font = Font(color="444444", name="Calibri")
        ws.cell(row=row, column=2, value=val).font = Font(bold=True, color="000000", name="Calibri")
        row += 1
    row += 1

    # 4 — CATEGORY TOTALS (moved above the summary + line items, R2.4)
    ws.cell(row=row, column=1, value="CATEGORY TOTALS").font = Font(bold=True, color="388BFD", name="Calibri")
    row += 1
    for cat, total, is_opt in ctx["category_totals"]:
        ws.cell(row=row, column=1, value=cat).font = Font(
            color=("F85149" if is_opt else "444444"), name="Calibri", bold=is_opt)
        cell = ws.cell(row=row, column=9, value=total)
        cell.font = Font(bold=True, color=("F85149" if is_opt else "000000"), name="Calibri")
        cell.number_format = "#,##0.00"
        cell.alignment = Alignment(horizontal="right")
        row += 1
    row += 1

    # 5 — PRICE SUMMARY (replicates the costings page; no cost per m²)
    ws.cell(row=row, column=1, value="PRICE SUMMARY").font = Font(bold=True, color="388BFD", name="Calibri")
    row += 1
    for pr in ctx["price_rows"]:
        is_total = pr["kind"] == "total"
        ws.merge_cells(f"A{row}:H{row}")
        lc = ws.cell(row=row, column=1, value=pr["label"])
        lc.font = Font(
            bold=True,
            color="E6EDF3" if is_total else "C9D1D9",
            size=12 if is_total else 11,
            name="Calibri",
        )
        lc.fill = grand_fill if is_total else total_fill
        lc.alignment = Alignment(horizontal="right", vertical="center")
        vc = ws.cell(row=row, column=9, value=pr["amount"])
        vc.font = Font(
            bold=True,
            color="FFFFFF" if is_total else "58A6FF",
            size=13 if is_total else 11,
            name="Calibri",
        )
        vc.fill = grand_fill if is_total else total_fill
        vc.number_format = "#,##0.00"
        vc.alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row].height = 22 if is_total else 18
        row += 1

    if ctx.get("zero_rule_note"):
        ws.merge_cells(f"A{row}:I{row}")
        zc = ws.cell(row=row, column=1, value=ctx["zero_rule_note"])
        zc.font = Font(bold=True, color="E02424", name="Calibri", size=11)
        zc.alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row].height = 18
        row += 1

    # 6 — line items, only when "with line items" was picked
    sections: list[dict] = []
    if ctx["include_items"]:
        row += 1
        cols = ["Category", "Material", "SAP Code", "Formula", "Quantity", "Unit",
                "Unit Price (R)", "Waste %", "Line Cost (R)"]
        for c, col in enumerate(cols, 1):
            cell = ws.cell(row=row, column=c, value=col)
            cell.font = Font(bold=True, color="E6EDF3", name="Calibri")
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal="center" if c > 4 else "left", vertical="center")
            cell.border = thin
        ws.row_dimensions[row].height = 18
        row += 1

        optional_cats = ctx["optional_cats"]
        current_cat = None
        current_section = None
        for item in ctx["items"]:
            cat = item["category"]
            if cat != current_cat:
                if current_section is not None:
                    current_section["last"] = row - 1
                    sections.append(current_section)
                ws.merge_cells(f"A{row}:H{row}")   # column I stays free for the section total
                _is_opt_cat = cat in optional_cats
                c = ws.cell(row=row, column=1, value=cat.upper())
                c.font = Font(bold=True,
                              color=("F85149" if _is_opt_cat else "58A6FF"),
                              name="Calibri")
                c.fill = cat_fill
                c.alignment = Alignment(horizontal="left", vertical="center")
                ws.row_dimensions[row].height = 16
                current_section = {"header": row, "first": row + 1, "last": None,
                                   "kind": "body", "cat": cat, "optional": _is_opt_cat}
                row += 1
                current_cat = cat

            mat_name = item["material"]
            if highlight and mat_name in override_materials:
                price_colour = "CC0000"
                row_tint = "FFF5F5"
            elif highlight and mat_name in recently_updated_mats:
                price_colour = "1A6FBF"
                row_tint = "F0F6FF"
            else:
                price_colour = "000000"
                row_tint = None

            cells_data = [
                ("", "left"), (item["material"], "left"), (item["material_code"], "left"),
                (item["formula"], "left"),
                (item["quantity"], "right"), (item["unit"], "center"),
                (item["unit_price"], "right"), (item["waste_pct"], "right"),
                (item["line_cost"], "right"),
            ]
            for c, (val, align) in enumerate(cells_data, 1):
                cell = ws.cell(row=row, column=c, value=val)
                if c == 7:
                    cell.font = Font(color=price_colour, name="Calibri", size=10,
                                     bold=(price_colour != "000000"))
                else:
                    cell.font = Font(color="000000", name="Calibri", size=10)
                if row_tint and c != 7:
                    cell.fill = PatternFill("solid", fgColor=row_tint)
                cell.alignment = Alignment(horizontal=align, vertical="center")
                cell.border = thin
                if c in (7, 9) and isinstance(val, (int, float)):
                    cell.number_format = "#,##0.00"
                if c == 8 and isinstance(val, (int, float)) and val:
                    cell.value = f"{val}%"
            ws.row_dimensions[row].height = 15
            row += 1

        if current_section is not None:
            current_section["last"] = row - 1
            sections.append(current_section)

        chassis = ctx["chassis"] or {}
        if chassis.get("items"):
            ws.merge_cells(f"A{row}:I{row}")
            c = ws.cell(
                row=row,
                column=1,
                value=(
                    f"CHASSIS  ({chassis.get('axle_count')}-axle · "
                    f"{chassis.get('tyre_style')} · "
                    f"{chassis.get('tyre_count')} tyres · {chassis.get('length')} m)"
                ),
            )
            c.font = Font(bold=True, color="58A6FF", name="Calibri")
            c.fill = cat_fill
            c.alignment = Alignment(horizontal="left", vertical="center")
            ws.row_dimensions[row].height = 16
            ch_header = row
            row += 1
            ch_first = row
            for it in chassis["items"]:
                cells_data = [
                    (it.get("kind", ""), "left"), (it.get("label", ""), "left"), ("", "left"),
                    ("", "left"),
                    (it.get("qty", 0), "right"), ("ea", "center"),
                    (it.get("unit_price", 0), "right"), ("", "right"),
                    (it.get("line_cost", 0), "right"),
                ]
                for cn, (val, align) in enumerate(cells_data, 1):
                    cell = ws.cell(row=row, column=cn, value=val)
                    cell.font = Font(color="000000", name="Calibri", size=10)
                    cell.alignment = Alignment(horizontal=align, vertical="center")
                    cell.border = thin
                    if cn in (5, 7, 9) and isinstance(val, (int, float)):
                        cell.number_format = "#,##0.00"
                ws.row_dimensions[row].height = 15
                row += 1
            if row > ch_first:
                sections.append({"header": ch_header, "first": ch_first, "last": row - 1, "kind": "chassis"})
            ws.merge_cells(f"A{row}:H{row}")
            lc = ws.cell(row=row, column=1, value="CHASSIS SUBTOTAL")
            lc.font = Font(bold=True, color="C9D1D9", name="Calibri")
            lc.fill = total_fill
            lc.alignment = Alignment(horizontal="right", vertical="center")
            vc = ws.cell(row=row, column=9, value=chassis.get("subtotal", 0))
            vc.font = Font(bold=True, color="58A6FF", name="Calibri")
            vc.fill = total_fill
            vc.number_format = "#,##0.00"
            vc.alignment = Alignment(horizontal="right", vertical="center")
            ws.row_dimensions[row].height = 18
            row += 1

    widths = [14, 42, 18, 32, 12, 8, 14, 9, 14]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    if highlight and (override_materials or recently_updated_mats):
        wl = wb.create_sheet("Legend")
        wl.column_dimensions["A"].width = 18
        wl.column_dimensions["B"].width = 52
        wl.column_dimensions["C"].width = 22

        wl.merge_cells("A1:C1")
        lt = wl["A1"]
        lt.value = "PRICE HIGHLIGHT LEGEND"
        lt.font = Font(bold=True, size=13, color="58A6FF", name="Calibri")
        lt.fill = PatternFill("solid", fgColor="0D1117")
        lt.alignment = Alignment(horizontal="center", vertical="center")
        wl.row_dimensions[1].height = 24

        for col, (lbl, clr) in enumerate([("Sample Price", "58A6FF"), ("Meaning", "8B949E"), ("Applies to", "8B949E")], 1):
            hc = wl.cell(row=2, column=col, value=lbl)
            hc.font = Font(bold=True, color=clr, name="Calibri")
            hc.fill = PatternFill("solid", fgColor="1C2333")
            hc.border = thin
            hc.alignment = Alignment(horizontal="center")
        wl.row_dimensions[2].height = 16

        leg_data = []
        if recently_updated_mats:
            leg_data.append((
                "1A6FBF", "F0F6FF", "R 123.45",
                "Price permanently updated in the material database (within last 7 days)",
                ", ".join(sorted(recently_updated_mats)[:5]) +
                (f" + {len(recently_updated_mats)-5} more" if len(recently_updated_mats) > 5 else ""),
            ))
        if override_materials:
            leg_data.append((
                "CC0000", "FFF5F5", "R 99.00",
                "Quote-only price override — not saved to database, applies to this quote only",
                ", ".join(sorted(override_materials)[:5]) +
                (f" + {len(override_materials)-5} more" if len(override_materials) > 5 else ""),
            ))

        for i, (fc, bg, sample, meaning, mats) in enumerate(leg_data, 3):
            wl.cell(row=i, column=1, value=sample).font = Font(bold=True, color=fc, name="Calibri", size=11)
            wl.cell(row=i, column=1).fill = PatternFill("solid", fgColor=bg)
            wl.cell(row=i, column=1).border = thin
            wl.cell(row=i, column=1).alignment = Alignment(horizontal="center", vertical="center")
            wl.cell(row=i, column=2, value=meaning).font = Font(color="111111", name="Calibri", size=10)
            wl.cell(row=i, column=2).fill = PatternFill("solid", fgColor=bg)
            wl.cell(row=i, column=2).border = thin
            wl.cell(row=i, column=2).alignment = Alignment(wrap_text=True, vertical="center")
            wl.cell(row=i, column=3, value=mats).font = Font(color="555555", name="Calibri", size=9, italic=True)
            wl.cell(row=i, column=3).fill = PatternFill("solid", fgColor=bg)
            wl.cell(row=i, column=3).border = thin
            wl.cell(row=i, column=3).alignment = Alignment(wrap_text=True, vertical="center")
            wl.row_dimensions[i].height = 36

        ws["A2"].value = (ws["A2"].value or "") + "  |  ⬤ Highlighted: price changes colour-coded  →  see Legend tab"

    # Collapsible row grouping (outline) — opens EXPANDED. Section-header totals
    # are always literal numbers (=SUM carries no cached value in openpyxl
    # output — Michael 4 Aug; byte-stability of the saved export is waived this
    # release, so the literal applies to previews AND approved exports).
    if sections:
        if ws.sheet_properties.outlinePr is None:
            from openpyxl.worksheet.properties import Outline
            ws.sheet_properties.outlinePr = Outline()
        ws.sheet_properties.outlinePr.summaryBelow = False
        ws.sheet_properties.outlinePr.applyStyles = False
        ws.sheet_view.showOutlineSymbols = True
        section_totals = ctx["section_totals"]
        for sec in sections:
            first, last, header = sec["first"], sec.get("last"), sec["header"]
            empty = last is None or last < first
            if not empty:
                for r in range(first, last + 1):
                    ws.row_dimensions[r].outline_level = 1   # detail rows only; never hidden
            if sec.get("kind") == "body":
                tot_value = 0 if empty else round(float(section_totals.get(sec.get("cat"), 0.0)), 2)
                tot = ws.cell(row=header, column=9, value=tot_value)
                tot.font = Font(bold=True,
                                color=("F85149" if sec.get("optional") else "58A6FF"),
                                name="Calibri")
                tot.fill = cat_fill
                tot.number_format = "#,##0.00"
                tot.alignment = Alignment(horizontal="right", vertical="center")
                tot.border = thin

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ── Word renderer ─────────────────────────────────────────────────────────────

def _render_docx(ctx: dict):
    """Render the shared doc ctx as a .docx (BytesIO) via python-docx (already a
    declared dependency — WO v4.33). Landscape when line items are included so
    the 8-column tables breathe; portrait for the compact totals-only variant."""
    import io

    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    ACCENT = RGBColor(0x1A, 0x6F, 0xBF)
    RED = RGBColor(0xE0, 0x24, 0x24)
    DIM = RGBColor(0x55, 0x55, 0x55)

    doc = Document()
    sec = doc.sections[0]
    if ctx["include_items"]:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    def para(text="", *, bold=False, size=None, color=None, center=False, space_after=4):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    def money(v):
        try:
            return f"{float(v):,.2f}"
        except (TypeError, ValueError):
            return "" if v in (None, "") else str(v)

    def cell_text(cell, text, *, bold=False, color=None, size=9, right=False):
        p = cell.paragraphs[0]
        if right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("" if text is None else str(text))
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color

    para(ctx["heading"], bold=True, size=16, center=True,
         color=(RED if ctx["is_repair"] else ACCENT))
    para(ctx["sub"], size=9, color=DIM, center=True)
    para(f"Client:  {ctx['client_name']}", bold=True, size=11, center=True, space_after=10)

    para("DIMENSIONS & BODY OPTIONS", bold=True, size=10, color=ACCENT)
    spec_rows = list(ctx["spec_pairs"]) + list(ctx["spec_options"])
    if spec_rows:
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for lbl, val in spec_rows:
            cells = t.add_row().cells
            cell_text(cells[0], lbl, color=DIM)
            cell_text(cells[1], val, bold=True)
    para(space_after=6)

    para("CATEGORY TOTALS", bold=True, size=10, color=ACCENT)
    if ctx["category_totals"]:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        cell_text(hdr[0], "Category", bold=True)
        cell_text(hdr[1], "Subtotal (R)", bold=True, right=True)
        for cat, total, is_opt in ctx["category_totals"]:
            cells = t.add_row().cells
            cell_text(cells[0], cat, bold=is_opt, color=(RED if is_opt else None))
            cell_text(cells[1], money(total), bold=True, right=True,
                      color=(RED if is_opt else None))
    para(space_after=6)

    para("PRICE SUMMARY", bold=True, size=10, color=ACCENT)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for pr in ctx["price_rows"]:
        cells = t.add_row().cells
        is_total = pr["kind"] == "total"
        cell_text(cells[0], pr["label"], bold=True,
                  size=11 if is_total else 10, color=(ACCENT if is_total else None))
        cell_text(cells[1], f"R {money(pr['amount'])}", bold=True, right=True,
                  size=11 if is_total else 10, color=(ACCENT if is_total else None))

    if ctx.get("zero_rule_note"):
        _zp = para(ctx["zero_rule_note"], bold=True, size=10, color=RED, space_after=4)
        _zp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if ctx["include_items"]:
        para(space_after=8)
        para("LINE ITEMS", bold=True, size=10, color=ACCENT)
        groups: list[tuple[str, list]] = []
        for it in ctx["items"]:
            cat = it.get("category") or "Uncategorised"
            if not groups or groups[-1][0] != cat:
                groups.append((cat, []))
            groups[-1][1].append(it)
        section_totals = ctx["section_totals"]
        for cat, items in groups:
            is_opt = cat in ctx["optional_cats"]
            para(cat.upper(), bold=True, size=9, space_after=2,
                 color=(RED if is_opt else ACCENT))
            t = doc.add_table(rows=1, cols=8)
            t.style = "Table Grid"
            headers = ["Material", "SAP Code", "Formula", "Qty", "Unit",
                       "Unit Price (R)", "Waste %", "Line Cost (R)"]
            for i, h in enumerate(headers):
                cell_text(t.rows[0].cells[i], h, bold=True, size=8,
                          right=(i in (3, 5, 6, 7)))
            for it in items:
                cells = t.add_row().cells
                waste = it.get("waste_pct")
                cell_text(cells[0], it.get("material") or "", size=8)
                cell_text(cells[1], it.get("material_code") or "", size=8)
                cell_text(cells[2], it.get("formula") or "", size=8)
                cell_text(cells[3], money(it.get("quantity")), size=8, right=True)
                cell_text(cells[4], it.get("unit") or "", size=8)
                cell_text(cells[5], money(it.get("unit_price")), size=8, right=True)
                cell_text(cells[6], f"{waste}%" if waste else "", size=8, right=True)
                cell_text(cells[7], money(it.get("line_cost")), size=8, right=True, bold=True)
            trow = t.add_row().cells
            cell_text(trow[6], f"{cat} total", size=8, right=True, color=DIM)
            cell_text(trow[7], money(section_totals.get(cat, 0.0)), size=8,
                      right=True, bold=True, color=(RED if is_opt else ACCENT))
            para(space_after=4)

        chassis = ctx["chassis"] or {}
        if chassis.get("items"):
            para(
                f"CHASSIS ({chassis.get('axle_count')}-axle · {chassis.get('tyre_style')} · "
                f"{chassis.get('tyre_count')} tyres · {chassis.get('length')} m)",
                bold=True, size=9, color=ACCENT, space_after=2)
            t = doc.add_table(rows=0, cols=4)
            t.style = "Table Grid"
            for it in chassis["items"]:
                cells = t.add_row().cells
                cell_text(cells[0], it.get("kind") or "", size=8)
                cell_text(cells[1], it.get("label") or "", size=8)
                cell_text(cells[2], money(it.get("qty")), size=8, right=True)
                cell_text(cells[3], money(it.get("line_cost")), size=8, right=True, bold=True)
            crow = t.add_row().cells
            cell_text(crow[2], "Chassis subtotal", size=8, right=True, color=DIM)
            cell_text(crow[3], money(chassis.get("subtotal")), size=8, right=True,
                      bold=True, color=ACCENT)

    if ctx.get("generated_at"):
        para(f"Generated {ctx['generated_at']}", size=8, color=DIM, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF renderer (ReportLab) ──────────────────────────────────────────────────

def _render_pdf(ctx: dict) -> bytes:
    """Render the shared doc ctx as a PDF (bytes) via ReportLab (A4 landscape,
    numbered footer). Same document order as the xlsx/docx renderers.

    Page 1 is a COVER: heading, client, dimensions & body options, category
    totals and the price summary (Michael 8 Aug — the summary used to flow onto
    page 2 whenever the category list was long, and on a short one the line
    items crowded in underneath it). Two things hold that shape:
      * the cover block is compact — body options pair up two per row (the same
        multi-column language the dimensions row above already uses) and the
        paddings are tight — so the whole cover fits one page for a realistic
        body (measured: ~14 categories at A4 landscape);
      * the line items always start on a fresh page.
    A KeepTogether around the totals+summary pair was tried and rejected: past
    the fitting limit it relocates BOTH tables to page 2 and leaves page 1
    nearly empty. Letting the category table split naturally is the better
    failure mode — the summary still follows immediately after it.
    """
    from xml.sax.saxutils import escape

    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.pdfgen import canvas
    from reportlab.platypus import (
        PageBreak, SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    def _fmt2(v):
        try:
            return "{:,.2f}".format(float(v))
        except (TypeError, ValueError):
            return "" if v in (None, "") else str(v)

    page_w, page_h = landscape(A4)
    left_margin = right_margin = 10 * mm
    avail_w = page_w - left_margin - right_margin

    optional_cats = set(ctx["optional_cats"])
    BLUE = colors.HexColor("#58A6FF")
    RED = colors.HexColor("#F85149")

    body = ParagraphStyle("cell", fontName="Helvetica", fontSize=8, leading=9.5)
    title_style = ParagraphStyle(
        "title", fontName="Helvetica-Bold", fontSize=14, leading=18,
        alignment=TA_CENTER,
        textColor=(colors.HexColor("#E02424") if ctx["is_repair"] else BLUE),
    )
    sub_style = ParagraphStyle(
        "sub", fontName="Helvetica", fontSize=9.5, leading=12,
        alignment=TA_CENTER, textColor=colors.HexColor("#C9D1D9"),
    )
    client_style = ParagraphStyle(
        "client", fontName="Helvetica-Bold", fontSize=11, leading=14,
        alignment=TA_CENTER, textColor=colors.HexColor("#0D1117"),
    )
    section_head = ParagraphStyle("sh", fontSize=10, leading=12)

    elements = []

    title_tbl = Table([[Paragraph(escape(ctx["heading"]), title_style)]], colWidths=[avail_w])
    title_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0D1117")),
        ("TOPPADDING", (0, 0), (-1, -1), 3.5 * mm),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5 * mm),
    ]))
    elements.append(title_tbl)

    sub_tbl = Table([[Paragraph(escape(ctx["sub"]), sub_style)]], colWidths=[avail_w])
    sub_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#161B22")),
        ("TOPPADDING", (0, 0), (-1, -1), 2 * mm),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2 * mm),
    ]))
    elements.append(sub_tbl)
    elements.append(Spacer(1, 2.5 * mm))

    elements.append(Paragraph("Client: " + escape(str(ctx["client_name"])), client_style))
    elements.append(Spacer(1, 2.5 * mm))

    # Dimensions + body options
    elements.append(Paragraph(
        '<font color="#1F3A5F"><b>DIMENSIONS &amp; BODY OPTIONS</b></font>', section_head))
    elements.append(Spacer(1, 1 * mm))
    if ctx["spec_pairs"]:
        drow = []
        for lbl, val in ctx["spec_pairs"]:
            drow += [str(lbl), "" if val in (None, "") else str(val)]
        lbl_w = avail_w * 0.14
        val_w = avail_w * 0.11
        dim_tbl = Table([drow], colWidths=[lbl_w, val_w] * len(ctx["spec_pairs"]))
        style = [("FONTSIZE", (0, 0), (-1, -1), 9),
                 ("TOPPADDING", (0, 0), (-1, -1), 0.8 * mm),
                 ("BOTTOMPADDING", (0, 0), (-1, -1), 0.8 * mm)]
        for i in range(len(ctx["spec_pairs"])):
            style.append(("TEXTCOLOR", (i * 2, 0), (i * 2, -1), colors.HexColor("#555555")))
            style.append(("FONTNAME", (i * 2 + 1, 0), (i * 2 + 1, -1), "Helvetica-Bold"))
        dim_tbl.setStyle(TableStyle(style))
        elements.append(dim_tbl)
    if ctx["spec_options"]:
        # Two option pairs per row — same multi-column shape as the dimensions
        # row above, and it halves the cover's tallest fixed block.
        opts = list(ctx["spec_options"])
        orows = []
        for i in range(0, len(opts), 2):
            chunk = opts[i:i + 2]
            row = []
            for lbl, val in chunk:
                row += [str(lbl), Paragraph(escape(str(val)), body)]
            if len(chunk) == 1:
                row += ["", ""]
            orows.append(row)
        opt_tbl = Table(orows, colWidths=[avail_w * 0.13, avail_w * 0.30] * 2)
        opt_tbl.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#555555")),
            ("TEXTCOLOR", (2, 0), (2, -1), colors.HexColor("#555555")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 0.5 * mm),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5 * mm),
        ]))
        elements.append(opt_tbl)
    elements.append(Spacer(1, 2.5 * mm))

    # Category totals — above the price summary and any line items (R2.4).
    # Built through a factory so the same block can be re-laid-out two-up when
    # a long category list would otherwise push the summary off page 1.
    def _category_totals_table(columns: int):
        head = ["Category", "Subtotal (R)"] * columns
        rows = [head]
        red_cells = []          # (row, col-block) pairs for optional sections
        entries = list(ctx["category_totals"])
        per_col = -(-len(entries) // columns)        # ceil
        cols = [entries[i * per_col:(i + 1) * per_col] for i in range(columns)]
        for r in range(per_col):
            row = []
            for c in range(columns):
                if r < len(cols[c]):
                    cat, total, is_opt = cols[c][r]
                    row += [str(cat), _fmt2(total)]
                    if is_opt:
                        red_cells.append((len(rows), c))
                else:
                    row += ["", ""]
            rows.append(row)
        widths = []
        for _ in range(columns):
            widths += [avail_w * (0.45 / columns), avail_w * (0.15 / columns)]
        tbl = Table(rows, colWidths=widths)
        style = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1C2333")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#E6EDF3")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("LINEBELOW", (0, 1), (-1, -1), 0.3, colors.HexColor("#D0D7DE")),
            ("TOPPADDING", (0, 0), (-1, -1), 0.7 * mm),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0.7 * mm),
        ]
        for c in range(columns):
            style.append(("ALIGN", (c * 2 + 1, 0), (c * 2 + 1, -1), "RIGHT"))
            style.append(("FONTNAME", (c * 2 + 1, 1), (c * 2 + 1, -1), "Helvetica-Bold"))
        for r, c in red_cells:
            style.append(("TEXTCOLOR", (c * 2, r), (c * 2 + 1, r), RED))
        tbl.setStyle(TableStyle(style))
        return tbl

    cover_tail: list = []
    if ctx["category_totals"]:
        cover_tail.append(_category_totals_table(1))
        cover_tail.append(Spacer(1, 2.5 * mm))

    # Price summary
    sdata = []
    grand_rows = []
    for i, pr in enumerate(ctx["price_rows"]):
        amount = pr["amount"]
        if pr["kind"] == "add":
            val = ("+ R " + _fmt2(amount)) if amount >= 0 else ("− R " + _fmt2(-amount))
        else:
            val = "R " + _fmt2(amount)
        sdata.append([pr["label"], val])
        if pr["kind"] == "total":
            grand_rows.append(i)
    summary_tbl = Table(sdata, colWidths=[avail_w * 0.75, avail_w * 0.25])
    s_style = [
        ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D7DE")),
        ("TOPPADDING", (0, 0), (-1, -1), 1.4 * mm),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1.4 * mm),
    ]
    for i in range(len(sdata)):
        if i in grand_rows:
            s_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#388BFD")))
            s_style.append(("TEXTCOLOR", (0, i), (-1, i), colors.white))
            s_style.append(("FONTSIZE", (0, i), (0, i), 11))
            s_style.append(("FONTSIZE", (1, i), (1, i), 12))
        else:
            s_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#0D4A8A")))
            s_style.append(("TEXTCOLOR", (0, i), (0, i), colors.HexColor("#C9D1D9")))
            s_style.append(("TEXTCOLOR", (1, i), (1, i), BLUE))
    summary_tbl.setStyle(TableStyle(s_style))
    cover_tail.append(summary_tbl)

    # The 3.2 m zero-rule notice belongs to the cover too — it explains the
    # totals directly above it, so it is measured and placed with them.
    if ctx.get("zero_rule_note"):
        cover_tail.append(Spacer(1, 2 * mm))
        cover_tail.append(Paragraph(
            escape(str(ctx["zero_rule_note"])),
            ParagraphStyle("zero_rule", fontName="Helvetica-Bold", fontSize=10,
                           leading=12, alignment=TA_RIGHT,
                           textColor=colors.HexColor("#E02424"))))

    # Does the whole cover fit on page 1? Measure rather than assume: a single
    # column of categories is the ratified look, so keep it whenever it fits
    # (every real body today — the widest is 11 categories) and only pair them
    # up two-per-row when it genuinely would not, which keeps even a
    # 25-category body's summary on page 1.
    def _stack_height(flowables) -> float:
        total = 0.0
        for f in flowables:
            try:
                total += f.wrap(avail_w, avail_h)[1]
            except Exception:
                return float("inf")     # unmeasurable → treat as overflowing
        return total

    avail_h = page_h - 12 * mm - 14 * mm         # matches the doc template below
    if ctx["category_totals"] and _stack_height(elements + cover_tail) > avail_h:
        cover_tail[0] = _category_totals_table(2)
    elements.extend(cover_tail)

    # Line items (grouped by category) — only when requested, and always from a
    # fresh page so page 1 stays the cover.
    if ctx["include_items"] and ctx["items"]:
        elements.append(PageBreak())
        bom_pct = [0.11, 0.26, 0.09, 0.19, 0.07, 0.05, 0.08, 0.06, 0.09]
        bom_w = [avail_w * p for p in bom_pct]
        header = ["Category", "Material", "SAP Code", "Formula", "Quantity",
                  "Unit", "Unit Price (R)", "Waste %", "Line Cost (R)"]
        data = [header]
        cat_rows, optional_rows = [], []
        current = object()
        for it in ctx["items"]:
            cat = it.get("category")
            if cat != current:
                data.append([str(cat or ""), "", "", "", "", "", "", "", ""])
                r = len(data) - 1
                cat_rows.append(r)
                if cat in optional_cats:
                    optional_rows.append(r)
                current = cat
            waste = it.get("waste_pct")
            data.append([
                "",
                Paragraph(escape(str(it.get("material") or "")), body),
                it.get("material_code") or "",
                Paragraph(escape(str(it.get("formula") or "")), body),
                _fmt2(it.get("quantity")),
                it.get("unit") or "",
                _fmt2(it.get("unit_price")),
                ("%s%%" % waste) if waste else "",
                _fmt2(it.get("line_cost")),
            ])

        bom_style = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1C2333")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#E6EDF3")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 8.5),
            ("FONTSIZE", (0, 1), (-1, -1), 8.5),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D7DE")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (4, 0), (4, -1), "RIGHT"),
            ("ALIGN", (5, 0), (5, -1), "CENTER"),
            ("ALIGN", (6, 0), (8, -1), "RIGHT"),
            ("TOPPADDING", (0, 0), (-1, -1), 1 * mm),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1 * mm),
        ]
        for r in cat_rows:
            bom_style.append(("SPAN", (0, r), (-1, r)))
            bom_style.append(("BACKGROUND", (0, r), (-1, r), colors.HexColor("#1F3A5F")))
            bom_style.append(("TEXTCOLOR", (0, r), (-1, r), RED if r in optional_rows else BLUE))
            bom_style.append(("FONTNAME", (0, r), (-1, r), "Helvetica-Bold"))
            bom_style.append(("ALIGN", (0, r), (-1, r), "LEFT"))
        bom_tbl = Table(data, colWidths=bom_w, repeatRows=1)
        bom_tbl.setStyle(TableStyle(bom_style))
        elements.append(bom_tbl)

        chassis = ctx["chassis"] or {}
        if chassis.get("items"):
            chead = "CHASSIS — %s-axle · %s · %s tyres · %s m" % (
                chassis.get("axle_count"), chassis.get("tyre_style"),
                chassis.get("tyre_count"), chassis.get("length"))
            cdata = [[chead, "", "", "", "", "", "", "", ""]]
            for it in chassis["items"]:
                cdata.append([
                    it.get("kind") or "",
                    Paragraph(escape(str(it.get("label") or "")), body),
                    "", "",
                    _fmt2(it.get("qty")), "ea", _fmt2(it.get("unit_price")),
                    "", _fmt2(it.get("line_cost")),
                ])
            cdata.append(["Chassis Subtotal", "", "", "", "", "", "", "",
                          _fmt2(chassis.get("subtotal"))])
            last = len(cdata) - 1
            ctbl = Table(cdata, colWidths=bom_w)
            ctbl.setStyle(TableStyle([
                ("SPAN", (0, 0), (-1, 0)),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3A5F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), BLUE),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("SPAN", (0, last), (-2, last)),
                ("ALIGN", (0, last), (-2, last), "RIGHT"),
                ("BACKGROUND", (0, last), (-1, last), colors.HexColor("#EFF3F8")),
                ("FONTNAME", (0, last), (-1, last), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D7DE")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("ALIGN", (4, 1), (4, last - 1), "RIGHT"),
                ("ALIGN", (5, 1), (5, last - 1), "CENTER"),
                ("ALIGN", (6, 1), (8, last), "RIGHT"),
                ("TOPPADDING", (0, 0), (-1, -1), 1 * mm),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1 * mm),
            ]))
            elements.append(Spacer(1, 4 * mm))
            elements.append(ctbl)

    gen_text = "Generated " + str(ctx.get("generated_at") or "")
    foot_left_x = left_margin
    foot_right_x = page_w - right_margin

    class _NumberedCanvas(canvas.Canvas):
        def __init__(self, *a, **k):
            canvas.Canvas.__init__(self, *a, **k)
            self._saved_pages = []

        def showPage(self):
            self._saved_pages.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._saved_pages)
            for state in self._saved_pages:
                self.__dict__.update(state)
                self.setFont("Helvetica", 8)
                self.setFillColor(colors.HexColor("#555555"))
                self.drawString(foot_left_x, 8 * mm, gen_text)
                self.drawRightString(
                    foot_right_x, 8 * mm,
                    "Page %d of %d" % (self._pageNumber, total))
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(A4),
        leftMargin=left_margin, rightMargin=right_margin,
        topMargin=12 * mm, bottomMargin=14 * mm,
        title=ctx["heading"],
    )
    doc.build(elements, canvasmaker=_NumberedCanvas)
    return buf.getvalue()


# ── Context assembly ──────────────────────────────────────────────────────────

def _spec_pairs(dims: dict) -> list:
    return [
        ("Length (m)", dims.get("length")),
        ("Width (m)", dims.get("width")),
        ("Height (m)", dims.get("height")),
    ]


def _spec_options_from_derived(derived) -> list:
    spec_options: list[tuple[str, str]] = []
    if derived:
        rd = derived.get("rear_door")
        if rd:
            spec_options.append((
                "DOOR TYPE",
                f"{rd['door_type']} — {rd['insulation']} ({rd['thickness_m']:.3f} m)"))
        for p in derived.get("panels") or []:
            label = "FLOOR INSULATION" if p["location"] == "FLOOR" else p["location"]
            spec_options.append(
                (label, f"{p['insulation']} ({p['thickness_m']:.3f} m)"))
        if derived.get("floor_type"):
            spec_options.append(("FLOOR TYPE", derived["floor_type"]))
    return spec_options


def _highlight_sets(result: dict, highlight: bool):
    override_materials: set = set()
    recently_updated_mats: set = set()
    if highlight:
        override_materials = set((result.get("overrides_by_name") or {}).keys())
        cutoff = datetime.now(timezone.utc) - timedelta(days=7)
        for item in result.get("items", []):
            lu = item.get("last_updated")
            if lu and item["material"] not in override_materials:
                try:
                    lu_dt = datetime.fromisoformat(lu)
                    if lu_dt.tzinfo is None:
                        lu_dt = lu_dt.replace(tzinfo=timezone.utc)
                    if lu_dt >= cutoff:
                        recently_updated_mats.add(item["material"])
                except Exception:
                    pass
    return override_materials, recently_updated_mats


def _default_ratios(result: dict) -> list[float]:
    """The page's currently-saved/selected ratio, as the one-element default."""
    rv = result.get("ratio_value")
    return parse_ratios([rv]) if rv else []


def _doc_ctx_for_record(rec: CalculationRecord, db: Session, *, detail, ratios_raw,
                        highlight: bool = False):
    """(doc_ctx, filename_stem) for an APPROVED/saved costing record."""
    dims = json.loads(rec.dimensions_json or "{}")
    result = json.loads(rec.result_json or "{}")
    result = strip_excluded_items(result)   # only selected items on exports

    tt = db.query(TrailerType).filter_by(id=rec.trailer_type_id).first()
    trailer_name = tt.name if tt else "Trailer"
    detail = detail if detail in VALID_DETAILS else "items"
    ratios = parse_ratios(ratios_raw) if ratios_raw is not None else _default_ratios(result)

    # Body-options block: same derivation as GET /api/calculations/{id} (the
    # costing-detail panel) — saved input_state + current BOM master data.
    from .calculator import _derive_body_options_display
    bom_rows = (db.query(BillOfMaterial)
                .filter_by(trailer_type_id=rec.trailer_type_id, is_body_option=True)
                .options(*_bom_load_options())
                .all())
    derived = _derive_body_options_display(
        bom_rows, result.get("input_state") or {}, result.get("body_variables"))

    quote_no = rec.quote_number or f"#{rec.id}"
    heading = f"{quote_no} — {body_type_with_length(trailer_name, dims.get('length'))}"
    is_repair = bool(rec.is_repair)
    sub = f"{trailer_name}  |  Report #{rec.id}  |  {rec.created_at.strftime('%d %B %Y') if rec.created_at else ''}"
    if is_repair:
        sub += "  |  REPAIR QUOTE"

    override_materials, recently_updated_mats = _highlight_sets(result, highlight)

    ctx = build_doc_ctx(
        mode="approved",
        heading=heading,
        sub=sub,
        client_name=(rec.customer.name if rec.customer else ""),
        spec_pairs=_spec_pairs(dims),
        spec_options=_spec_options_from_derived(derived),
        result=result,
        ratios=ratios,
        detail=detail,
        db=db,
        heading_color="E02424" if is_repair else "58A6FF",
        is_repair=is_repair,
        highlight=highlight,
        override_materials=override_materials,
        recently_updated_mats=recently_updated_mats,
        generated_at=datetime.now().strftime("%d %b %Y %H:%M"),
    )
    username = rec.user.username if rec.user else "unknown"
    stem = f"Costing_{trailer_name.replace(' ', '_')}_{rec.id}_{username}"
    return ctx, stem


def _doc_ctx_for_preview(body: dict, db: Session):
    """(doc_ctx, filename_stem) for the LIVE, not-yet-approved calculator result.
    Nothing is written to the DB and no quote number is consumed."""
    result = body.get("result")
    items = result.get("items") if isinstance(result, dict) else None
    if not isinstance(items, list) or not items:
        raise HTTPException(status_code=400,
                            detail="No calculated result to preview — calculate first")
    dims = body.get("dims")
    if not isinstance(dims, dict):
        dims = {}

    # Prefer the DB name (works for inactive body types too — edit sessions);
    # the client-supplied label is only a fallback for the heading.
    trailer_name = ""
    tt = None
    tt_id = body.get("trailer_type_id")
    if tt_id is not None:
        try:
            tt = db.query(TrailerType).filter_by(id=int(tt_id)).first()
        except (TypeError, ValueError):
            tt = None
        if tt:
            trailer_name = tt.name
    if not trailer_name:
        trailer_name = str(body.get("trailer_name") or "").strip() or "Body Type"

    result = strip_excluded_items(result)  # match saved-export semantics

    bom_rows = []
    if tt is not None:
        bom_rows = (db.query(BillOfMaterial)
                    .filter_by(trailer_type_id=tt.id).all())

    # Line items in the SAME order as the costings page (Michael 4 Aug). Sheet
    # mode (default): items keyed by BOM sort_order; alpha mode: sections A–Z.
    items_live = list(result.get("items", []))
    if items_live:
        mode = str(body.get("bom_sort_mode") or "sheet").lower()
        if mode == "alpha":
            items_live.sort(key=lambda it: ((it.get("category") or "Uncategorised"),
                                            str(it.get("material") or "")))
        else:
            so_by_bom = {r.id: r.sort_order for r in bom_rows
                         if r.sort_order is not None}
            keys: dict[int, float] = {}
            for idx, it in enumerate(items_live):
                bid = it.get("bom_id")
                keys[id(it)] = so_by_bom.get(bid, idx) if bid is not None else idx
            cat_first: dict[str, float] = {}
            for it in items_live:
                cat = it.get("category") or "Uncategorised"
                k = keys[id(it)]
                if cat not in cat_first or k < cat_first[cat]:
                    cat_first[cat] = k
            items_live.sort(key=lambda it: (
                cat_first[it.get("category") or "Uncategorised"], keys[id(it)]))
        result = dict(result)
        result["items"] = items_live

    # Spec block: the SELECTED body options exactly like the calculator's Body
    # Options panel — canonical read-only decoder, from the client's live state.
    derived = None
    if bom_rows:
        from .calculator import _derive_body_options_display
        input_state = {
            "body_option_selections": body.get("body_option_selections") or {},
            "ui_snapshot": {"drd_srd": body.get("drd_srd") or {}},
        }
        derived = _derive_body_options_display(
            bom_rows, input_state, (result.get("body_variables") or {}))

    detail = body.get("detail") if body.get("detail") in VALID_DETAILS else "items"
    ratios = (parse_ratios(body.get("ratios"))
              if "ratios" in body else _default_ratios(result))

    today = datetime.now()
    ctx = build_doc_ctx(
        mode="preview",
        heading=f"Testing — {body_type_with_length(trailer_name, dims.get('length'))}",
        sub=f"{trailer_name}  |  {today.strftime('%d %B %Y')}",
        client_name=str(body.get("customer_name") or "").strip(),
        spec_pairs=_spec_pairs(dims),
        spec_options=_spec_options_from_derived(derived),
        result=result,
        ratios=ratios,
        detail=detail,
        db=db,
        generated_at=today.strftime("%d %b %Y %H:%M"),
    )
    safe_name = "".join(ch for ch in trailer_name if ch not in '\\/:*?"<>|\r\n').strip() or "Body Type"
    stem = f"Testing - {safe_name} - {today.strftime('%Y-%m-%d')}"
    return ctx, stem


# ── Preview endpoints (live result; DB untouched) ─────────────────────────────

async def _export_preview_impl(request: Request, db: Session, force_format: str | None):
    body = await request.json()
    fmt = force_format or str(body.get("format") or "excel").lower()
    if fmt not in VALID_FORMATS:
        raise HTTPException(status_code=400,
                            detail=f"format must be one of {', '.join(VALID_FORMATS)}")
    _require_export_user(request, db, fmt)
    ctx, stem = _doc_ctx_for_preview(body, db)
    buf = _render_document(fmt, ctx)
    ext, media = _FORMAT_MEDIA[fmt]
    return StreamingResponse(
        buf,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{stem}.{ext}"'},
    )


@router.post("/api/export/preview")
async def export_preview(request: Request, db: Session = Depends(get_db)):
    """Excel / Word / PDF of the LIVE calculator result (v1.44 R3–R5).

    Body: {result, dims, trailer_type_id, trailer_name, body_option_selections,
    drd_srd, bom_sort_mode, customer_name, format: excel|word|pdf,
    detail: totals|items, ratios: [divisors]}."""
    return await _export_preview_impl(request, db, force_format=None)


@router.post("/api/export/excel-preview")
async def export_excel_preview(request: Request, db: Session = Depends(get_db)):
    """v1.44 F1 route, kept as an alias of /api/export/preview with the format
    pinned to Excel (old clients / cached JS keep working)."""
    return await _export_preview_impl(request, db, force_format="excel")


# ── Emailing a costing document (v1.45) ───────────────────────────────────────
# One recipient, the sender Cc'd where their account carries an address, the
# rendered document attached in whatever format/detail/ratios the dialog was set
# to. Ratified with Michael 10 Aug: recipient pre-filled from the costing's
# Attention contact but freely typed over; PREVIEWS are for internal review, so
# their body says so in as many words; fixed subject + an optional free-text note.

_EMAIL_RE = re.compile(r"^[^@\s,;]+@[^@\s,;]+\.[^@\s,;]+$")


def _allowlist() -> tuple[set[str], set[str]]:
    """(domains, addresses) — both lower-cased. Read per call, not at import, so
    an env change takes effect on restart without a code path caring."""
    from ..config import settings
    def _split(raw):
        return {p.strip().lower() for p in str(raw or "").split(",") if p.strip()}
    return (_split(getattr(settings, "COSTING_EMAIL_ALLOWED_DOMAINS", "")),
            _split(getattr(settings, "COSTING_EMAIL_ALLOWED_ADDRESSES", "")))


def _is_internal(addr: str) -> bool:
    """v1.45.1 — a costing carries a customer's full pricing, so it may only be
    emailed to an internal mailbox (Michael 10 Aug, after a test send reached an
    outside domain). Deny-by-default and FAIL CLOSED: an empty allowlist allows
    nothing at all rather than everything.

    Domain match is on the FULL domain, exactly — not a suffix test. `endswith`
    would wave through `finance@noticecoldgrp.co.za` and
    `x@icecoldgrp.co.za.attacker.com`; sub-domains are deliberately not inherited.
    """
    addr = (addr or "").strip().lower()
    if "@" not in addr:
        return False
    domains, addresses = _allowlist()
    if addr in addresses:
        return True
    return addr.rsplit("@", 1)[1] in domains


def _allowlist_message() -> str:
    domains, addresses = _allowlist()
    allowed = sorted(f"@{d}" for d in domains) + sorted(addresses)
    return ("Costings may only be emailed to internal addresses "
            f"({', '.join(allowed) or 'none configured'}). "
            "To send one to a customer, download it and attach it yourself.")


def _validated_recipient(raw) -> str:
    to = str(raw or "").strip()
    if not to:
        raise HTTPException(status_code=400, detail="Enter an email address to send to.")
    if not _EMAIL_RE.match(to):
        raise HTTPException(status_code=400,
                            detail=f"“{to[:80]}” doesn’t look like an email address.")
    if not _is_internal(to):
        # 403, not 400: the address is well-formed, it is simply not permitted.
        raise HTTPException(status_code=403, detail=_allowlist_message())
    return to


def _sender_cc(user) -> str | None:
    """The sender's own address, when their account has one AND it is internal.
    Many accounts still carry email='' (the column only landed in migration
    0030), so this is best-effort — never a reason to block the send. The
    allowlist applies here too: a user whose account carries an outside address
    must not become a side-channel for the document."""
    addr = (getattr(user, "email", "") or "").strip()
    if not addr or not _is_internal(addr):
        return None
    return addr


def _email_body(ctx: dict, *, mode: str, sender: str, note: str) -> tuple[str, str]:
    """(subject, body). Approved mail is quotation-shaped; preview mail states in
    the first line that it is an internal draft, so a forwarded copy can't be
    mistaken for a quote (it is headed "Testing — …" and has no quote number)."""
    heading = ctx["heading"]
    client = ctx["client_name"]
    if mode == "preview":
        subject = f"DRAFT (not a quotation) — {heading}"
        opening = ("This is an INTERNAL DRAFT costing for review — it is not a quotation, "
                   "it carries no quote number, and the figures may still change.")
    else:
        subject = f"Costing — {heading}"
        opening = "Please find the costing attached."
    lines = [opening, "", f"Body type:  {heading}", f"Client:     {client}"]
    totals = [r for r in ctx["price_rows"] if r["kind"] == "total"]
    if totals:
        lines.append("")
        for r in totals:
            lines.append(f"{r['label']}:  R {r['amount']:,.2f}")
    if note.strip():
        lines += ["", "-----", note.strip(), "-----"]
    lines += ["", f"Sent from the ICB MES by {sender}."]
    return subject, "\n".join(lines)


def _send_costing_email(*, fmt: str, ctx: dict, stem: str, mode: str,
                        user, to_raw, note_raw) -> dict:
    """Render + send. Shared by the preview and approved email endpoints."""
    from ..services.notifications import (EmailNotConfigured, EmailSendFailed,
                                          send_document_email)

    to = _validated_recipient(to_raw)
    note = str(note_raw or "")
    ext, media = _FORMAT_MEDIA[fmt]
    subtype = media.rsplit("/", 1)[-1]          # the long OOXML subtypes ride as-is
    blob = _render_document(fmt, ctx).getvalue()
    sender_name = getattr(user, "username", "") or "the ICB MES"
    subject, body = _email_body(ctx, mode=mode, sender=sender_name, note=note)
    cc = _sender_cc(user)
    try:
        send_document_email(subject=subject, body=body, to=to, cc=cc,
                            attachment=(f"{stem}.{ext}", blob, subtype))
    except EmailNotConfigured as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except EmailSendFailed as e:
        raise HTTPException(status_code=502,
                            detail=f"The email could not be sent — {e}") from e
    return {"ok": True, "to": to, "cc": cc, "filename": f"{stem}.{ext}",
            "subject": subject}


@router.get("/api/export/email-policy")
async def email_policy(request: Request, db: Session = Depends(get_db)):
    """The recipient allowlist, so a dialog can state the rule up front and avoid
    pre-filling an address the server will refuse. Read-only; the server remains
    the enforcement point — this is only there to keep the UI honest."""
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401)
    domains, addresses = _allowlist()
    return {"domains": sorted(domains), "addresses": sorted(addresses),
            "message": _allowlist_message()}


@router.post("/api/export/preview/email")
async def email_preview(request: Request, db: Session = Depends(get_db)):
    """Email the LIVE (not-yet-approved) costing document. Same body as
    /api/export/preview plus {to, note}. Writes NOTHING to the DB and consumes no
    quote number — emailing a preview is still a preview."""
    body = await request.json()
    fmt = str(body.get("format") or "excel").lower()
    if fmt not in VALID_FORMATS:
        raise HTTPException(status_code=400,
                            detail=f"format must be one of {', '.join(VALID_FORMATS)}")
    user = _require_export_user(request, db, fmt)
    ctx, stem = _doc_ctx_for_preview(body, db)
    return _send_costing_email(fmt=fmt, ctx=ctx, stem=stem, mode="preview",
                               user=user, to_raw=body.get("to"),
                               note_raw=body.get("note"))


@router.post("/results/{record_id}/export/email")
async def email_approved(record_id: int, request: Request,
                         db: Session = Depends(get_db)):
    """Email an APPROVED costing document. Body: {to, note, format, detail, ratios}."""
    body = await request.json()
    fmt = str(body.get("format") or "pdf").lower()
    if fmt not in VALID_FORMATS:
        raise HTTPException(status_code=400,
                            detail=f"format must be one of {', '.join(VALID_FORMATS)}")
    user = _require_export_user(request, db, fmt)
    rec = _get_record_or_404(record_id, db)
    ctx, stem = _doc_ctx_for_record(rec, db, detail=body.get("detail"),
                                    ratios_raw=body.get("ratios"))
    return _send_costing_email(fmt=fmt, ctx=ctx, stem=stem, mode="approved",
                               user=user, to_raw=body.get("to"),
                               note_raw=body.get("note"))


# ── Approved (saved-record) exports ───────────────────────────────────────────

def _get_record_or_404(record_id: int, db: Session) -> CalculationRecord:
    rec = db.query(CalculationRecord).filter_by(id=record_id).first()
    if not rec:
        raise HTTPException(status_code=404)
    return rec


@router.get("/results/{record_id}/export/excel")
async def export_excel(
    record_id: int,
    request: Request,
    db: Session = Depends(get_db),
    highlight: int = 0,
    detail: str | None = None,
    ratios: str | None = None,
):
    """Approved costing as Excel. highlight=1 → colour-code price changes."""
    _require_export_user(request, db, "excel")
    rec = _get_record_or_404(record_id, db)
    ctx, stem = _doc_ctx_for_record(rec, db, detail=detail, ratios_raw=ratios,
                                    highlight=bool(highlight))
    buf = _render_xlsx(ctx)
    return StreamingResponse(
        buf,
        media_type=_FORMAT_MEDIA["excel"][1],
        headers={"Content-Disposition": f'attachment; filename="{stem}.xlsx"'},
    )


@router.get("/results/{record_id}/export/word")
async def export_word(
    record_id: int,
    request: Request,
    db: Session = Depends(get_db),
    detail: str | None = None,
    ratios: str | None = None,
):
    """Approved costing as Word (v1.44 R4 — python-docx)."""
    _require_export_user(request, db, "word")
    rec = _get_record_or_404(record_id, db)
    try:
        ctx, stem = _doc_ctx_for_record(rec, db, detail=detail, ratios_raw=ratios)
        buf = _render_docx(ctx)
    except HTTPException:
        raise
    except Exception:
        import logging
        logging.getLogger(__name__).exception("export_word failed (record %s)", record_id)
        raise HTTPException(status_code=500, detail="Word generation failed")
    return StreamingResponse(
        buf,
        media_type=_FORMAT_MEDIA["word"][1],
        headers={"Content-Disposition": f'attachment; filename="{stem}.docx"'},
    )


@router.get("/results/{record_id}/export/pdf")
async def export_pdf(
    record_id: int,
    request: Request,
    db: Session = Depends(get_db),
    detail: str | None = None,
    ratios: str | None = None,
):
    _require_export_user(request, db, "pdf")
    rec = _get_record_or_404(record_id, db)
    try:
        ctx, stem = _doc_ctx_for_record(rec, db, detail=detail, ratios_raw=ratios)
        pdf_bytes = _render_pdf(ctx)
    except HTTPException:
        raise
    except Exception:
        # WO v4.36d §3.1 — log the full exception server-side; return a GENERIC
        # client message (the prior detail=f"...{exc}" leaked internals). 500 preserved.
        import logging
        logging.getLogger(__name__).exception("export_pdf failed (record %s)", record_id)
        raise HTTPException(status_code=500, detail="PDF generation failed")
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{stem}.pdf"'},
    )


@router.get("/results/{record_id}/report")
async def report_for_record(record_id: int, request: Request, db: Session = Depends(get_db)):
    """Render the report PDF using the trailer's resolved ReportTemplate."""
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401)
    if not user_can(user, "quote.generate", db):
        raise HTTPException(status_code=403, detail="Permission denied: quote.generate")

    rec = db.query(CalculationRecord).filter_by(id=record_id).first()
    if not rec:
        raise HTTPException(status_code=404, detail="Costing record not found")

    tt = db.query(TrailerType).filter_by(id=rec.trailer_type_id).first()
    tmpl = resolve_report_template(tt)
    if not tmpl:
        raise HTTPException(status_code=400, detail="No report template is assigned to this trailer type.")

    dims = json.loads(rec.dimensions_json or "{}")
    result = json.loads(rec.result_json or "{}")
    result = strip_excluded_items(result)  # only selected items on the report
    customer = None
    if rec.customer:
        customer = {
            "name": rec.customer.name or "",
            "email": rec.customer.email or "",
            "telephone": rec.customer.telephone or "",
            # Attention-of contact — the 0035 write-time snapshot, NOT a live join, so a
            # re-render years later still shows the person the quote was addressed to.
            "contact_name": rec.contact_name or "",
            "contact_email": rec.contact_email or "",
            "contact_telephone": rec.contact_telephone or "",
        }

    try:
        from ..report_engine import render_by_slug
        pdf_bytes = render_by_slug(
            slug=tmpl.slug,
            record_id=record_id,
            customer=customer,
            dimensions=dims,
            result=result,
            created_at=rec.created_at,
            quote_number=rec.quote_number,
        )
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Report generation failed: {e}")

    safe_customer = (
        "".join(ch for ch in (rec.customer.name if rec.customer else "Customer")
                if ch.isalnum() or ch in (" ", "_", "-"))
        .strip()
        .replace(" ", "_")
    ) or "Customer"
    safe_slug = tmpl.slug.replace("/", "_")
    filename = f"{safe_slug}_{record_id}_{safe_customer}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )


@router.get("/results/{record_id}/report/explosive-quote")
async def report_explosive_quote_compat(record_id: int, request: Request, db: Session = Depends(get_db)):
    return await report_for_record(record_id=record_id, request=request, db=db)
