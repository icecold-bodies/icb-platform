"""v1.44 preview/export formats — backend units for the six ratified requirements.

R1 — costings.price_master_edit: role 'full' saves permanent BOM prices via
     PUT /api/bom/{id} (price-only body) with a BomOverrideHistory audit row;
     'sales' cannot; admin unchanged; every other BOM field stays admin-only.
R2/R4 — one shared document context renders Excel / Word / PDF with identical
     content order; multi-ratio totals equal _apply_chassis_and_margin outputs
     for each ratio (the formula is reused, never reimplemented); the
     no-line-items variant contains zero item rows; headings + client line per
     mode (preview "Testing — …", approved "{quote|#id} — …").
R3 — the canonical ratio list (services.document_context.RATIO_OPTIONS) is
     parity-locked against BOTH calculator templates' #f-ratio selects.
R6 — GET /api/calculations rows carry body_length.
R7 — previews write nothing (covered in test_excel_preview_and_trailer_active
     for the alias; re-asserted here for the format-parameterised route).

Sessions are real UserSession rows via raw Cookie headers (banked pattern —
the routes use inline get_current_user/require_* chokepoints).
"""
import io
import json
import re
import uuid
from datetime import datetime
from pathlib import Path

import pytest

FIXED_RECORD_ID = 900145            # sibling file uses 900144
FIXED_CREATED_AT = datetime(2026, 2, 20, 9, 15, 0)
FIXED_TT_NAME = "V144 FORMATS BODY"

RESULT_FIXTURE = {
    "items": [
        {"category": "FLOOR", "material": "FMT FLOOR SHEET", "material_code": "FF-1",
         "formula": "L*W", "quantity": 10.0, "unit": "m2", "unit_price": 100.0,
         "waste_pct": 5, "line_cost": 1050.0, "last_updated": None},
        {"category": "SIDES", "material": "FMT SIDE PANEL", "material_code": "FS-1",
         "formula": "L*H*2", "quantity": 2.0, "unit": "ea", "unit_price": 375.0,
         "waste_pct": 0, "line_cost": 750.0, "last_updated": None},
    ],
    "category_totals": {"FLOOR": 1050.0, "SIDES": 750.0},
    "cost_per_sqm": 96.0,
    "grand_total": 1800.0,
    "profit_margin": 10,
    "profit_amount": 180.0,
    "ratio_value": 0.55,
    "ratio_label": "55%",
    "ratio_amount": 1600.0,
    "selling_price": 3600.0,
}
DIMS_FIXTURE = {"length": 13.6, "width": 2.5, "height": 2.6}


# ── fixtures ──────────────────────────────────────────────────────────────────
@pytest.fixture(scope="module")
def app_mod():
    import app.main as m
    from starlette.testclient import TestClient
    with TestClient(m.app) as _c:   # triggers startup → seeds permissions
        yield m


@pytest.fixture(scope="module")
def client(app_mod):
    from starlette.testclient import TestClient
    with TestClient(app_mod.app) as c:
        yield c


def _make_session(username: str) -> dict:
    from app.database import SessionLocal, User, UserSession
    sid = f"v144f-{uuid.uuid4().hex[:12]}"
    csrf = f"csrf-{sid}"
    with SessionLocal() as db:
        u = db.query(User).filter_by(username=username).first()
        assert u, f"user {username!r} missing"
        db.merge(UserSession(id=sid, user_id=u.id, role=u.role,
                             expires_at=None, csrf_token=csrf))
        db.commit()
    return {"Cookie": f"session_id={sid}", "X-CSRF-Token": csrf}


def _throwaway_user(role: str):
    from app.database import SessionLocal, User
    uname = f"t_{role}_{uuid.uuid4().hex[:6]}"
    with SessionLocal() as db:
        db.add(User(username=uname, password_hash="x", role=role))
        db.commit()
    return uname


def _drop_user(headers: dict, uname: str):
    from app.database import SessionLocal, User, UserSession
    with SessionLocal() as db:
        db.query(UserSession).filter_by(
            id=headers["Cookie"].split("session_id=")[1]).delete()
        db.query(User).filter_by(username=uname).delete()
        db.commit()


@pytest.fixture(scope="module")
def admin_headers(app_mod):
    return _make_session("admin")


@pytest.fixture(scope="module")
def full_headers(app_mod):
    uname = _throwaway_user("full")
    headers = _make_session(uname)
    yield headers
    _drop_user(headers, uname)


@pytest.fixture(scope="module")
def sales_headers(app_mod):
    uname = _throwaway_user("sales")
    headers = _make_session(uname)
    yield headers
    _drop_user(headers, uname)


@pytest.fixture(scope="module")
def planner_headers(app_mod):
    uname = _throwaway_user("planner")
    headers = _make_session(uname)
    yield headers
    _drop_user(headers, uname)


@pytest.fixture(scope="module")
def seeded(app_mod):
    """Deterministic trailer + record + one plain BOM row (price-save target)."""
    from app.database import (SessionLocal, TrailerType, CalculationRecord,
                              BillOfMaterial, BomOverrideHistory, Material, User)
    mat_name = f"{FIXED_TT_NAME} PRICE MAT"
    with SessionLocal() as db:
        admin = db.query(User).filter_by(username="admin").first()
        tt = db.query(TrailerType).filter_by(name=FIXED_TT_NAME).first()
        if not tt:
            tt = TrailerType(name=FIXED_TT_NAME, description="v1.44 formats fixture")
            db.add(tt)
            db.flush()
        mat = db.query(Material).filter_by(name=mat_name).first()
        if not mat:
            mat = Material(name=mat_name, unit_of_measure="ea", price_per_unit=50.0)
            db.add(mat)
            db.flush()
        bom = db.query(BillOfMaterial).filter_by(
            trailer_type_id=tt.id, material_id=mat.id).first()
        if not bom:
            bom = BillOfMaterial(trailer_type_id=tt.id, material_id=mat.id,
                                 formula_expression="1", waste_percentage=0,
                                 bom_section="FLOOR")
            db.add(bom)
            db.flush()
        rec = db.query(CalculationRecord).filter_by(id=FIXED_RECORD_ID).first()
        if not rec:
            rec = CalculationRecord(
                id=FIXED_RECORD_ID,
                trailer_type_id=tt.id, user_id=admin.id,
                dimensions_json=json.dumps(DIMS_FIXTURE),
                result_json=json.dumps(RESULT_FIXTURE),
                created_at=FIXED_CREATED_AT, status="pending", is_repair=False)
            db.add(rec)
        db.commit()
        ids = {"tt_id": tt.id, "rec_id": FIXED_RECORD_ID,
               "bom_id": bom.id, "mat_id": mat.id}
    yield ids
    with SessionLocal() as db:
        db.query(BomOverrideHistory).filter_by(bom_id=ids["bom_id"]).delete()
        db.query(CalculationRecord).filter_by(id=FIXED_RECORD_ID).delete()
        db.query(BillOfMaterial).filter_by(id=ids["bom_id"]).delete()
        db.query(Material).filter_by(id=ids["mat_id"]).delete()
        db.query(TrailerType).filter_by(id=ids["tt_id"]).delete()
        db.commit()


# ── helpers ───────────────────────────────────────────────────────────────────
def _sheet_cells(xlsx_bytes: bytes):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
    ws = wb.active
    col_a = [ws.cell(row=r, column=1).value for r in range(1, ws.max_row + 1)]
    col_i = [ws.cell(row=r, column=9).value for r in range(1, ws.max_row + 1)]
    flat = "|".join(str(c.value) for row in ws.iter_rows() for c in row
                    if c.value is not None)
    return ws, col_a, col_i, flat


def _docx_text(docx_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "|".join(parts)


def _expected_total(ratio: float) -> float:
    from app.routers.calculator import _apply_chassis_and_margin
    calc = _apply_chassis_and_margin(
        {"grand_total": RESULT_FIXTURE["grand_total"]},
        {"profit_margin": RESULT_FIXTURE["profit_margin"], "ratio_value": ratio},
        None)
    return calc["selling_price"]


# ── R3: ratio list single source ──────────────────────────────────────────────
def test_ratio_options_match_both_calculator_templates():
    from app.services.document_context import RATIO_OPTIONS
    tmpl_dir = Path(__file__).resolve().parents[1] / "app" / "templates"
    for tmpl in ("calculator.html", "calculator2.html"):
        html = (tmpl_dir / tmpl).read_text(encoding="utf-8")
        sel = re.search(r'<select[^>]*id="f-ratio"[^>]*>(.*?)</select>', html, re.S)
        assert sel, f"{tmpl}: #f-ratio select not found"
        opts = re.findall(r'<option value="([^"]+)">([^<]+)</option>', sel.group(1))
        assert [(float(v), lbl) for v, lbl in opts] == RATIO_OPTIONS, (
            f"{tmpl} #f-ratio drifted from document_context.RATIO_OPTIONS — "
            "update BOTH in the same change")


def test_permission_catalogue_seeded(app_mod):
    from app.database import SessionLocal, Permission, RolePermission
    with SessionLocal() as db:
        for perm_name in ("costings.price_master_edit", "export.word"):
            p = db.query(Permission).filter_by(name=perm_name).first()
            assert p, f"{perm_name} not seeded"
            grants = {r.role for r in db.query(RolePermission).filter_by(
                permission_id=p.id).all()}
            assert "full" in grants, f"{perm_name} lacks the 'full' grant"


# ── R1: Nadie's permanent price save ─────────────────────────────────────────
def _audit_rows(bom_id):
    from app.database import SessionLocal, BomOverrideHistory
    with SessionLocal() as db:
        return db.query(BomOverrideHistory).filter_by(bom_id=bom_id) \
                 .order_by(BomOverrideHistory.id).all()


def test_full_saves_price_with_audit_row(client, full_headers, seeded):
    bom_id = seeded["bom_id"]
    before = len(_audit_rows(bom_id))
    r = client.put(f"/api/bom/{bom_id}", headers=full_headers,
                   json={"unit_price_override": 222.75})
    assert r.status_code == 200, r.text
    rows = _audit_rows(bom_id)
    assert len(rows) == before + 1
    assert rows[-1].old_price is None and rows[-1].new_price == 222.75
    assert rows[-1].material_name and rows[-1].trailer_type_name
    assert rows[-1].batch_at is not None            # bulk-undo can revert it
    # clear ("Restore to base") — same permission, audited too
    r = client.put(f"/api/bom/{bom_id}", headers=full_headers,
                   json={"unit_price_override": None})
    assert r.status_code == 200, r.text
    rows = _audit_rows(bom_id)
    assert len(rows) == before + 2
    assert rows[-1].old_price == 222.75 and rows[-1].new_price is None


def test_no_audit_row_on_noop_save(client, full_headers, seeded):
    bom_id = seeded["bom_id"]
    client.put(f"/api/bom/{bom_id}", headers=full_headers,
               json={"unit_price_override": 100.0})
    before = len(_audit_rows(bom_id))
    r = client.put(f"/api/bom/{bom_id}", headers=full_headers,
                   json={"unit_price_override": 100.0})
    assert r.status_code == 200
    assert len(_audit_rows(bom_id)) == before       # unchanged value → no row
    client.put(f"/api/bom/{bom_id}", headers=full_headers,
               json={"unit_price_override": None})


def test_sales_cannot_save_price(client, sales_headers, seeded):
    r = client.put(f"/api/bom/{seeded['bom_id']}", headers=sales_headers,
                   json={"unit_price_override": 9.99})
    assert r.status_code == 403
    assert "costings.price_master_edit" in r.json()["detail"]


def test_full_cannot_touch_admin_only_fields(client, full_headers, seeded):
    r = client.put(f"/api/bom/{seeded['bom_id']}", headers=full_headers,
                   json={"formula_expression": "2"})
    assert r.status_code == 403
    assert r.json()["detail"] == "Admin access required"
    # …and a mixed body (price + admin field) is admin-only too.
    r = client.put(f"/api/bom/{seeded['bom_id']}", headers=full_headers,
                   json={"unit_price_override": 1.0, "formula_expression": "2"})
    assert r.status_code == 403


def test_admin_price_save_unchanged(client, admin_headers, seeded):
    bom_id = seeded["bom_id"]
    before = len(_audit_rows(bom_id))
    r = client.put(f"/api/bom/{bom_id}", headers=admin_headers,
                   json={"unit_price_override": 55.5})
    assert r.status_code == 200, r.text
    assert len(_audit_rows(bom_id)) == before + 1
    client.put(f"/api/bom/{bom_id}", headers=admin_headers,
               json={"unit_price_override": None})


def test_any_user_still_writes_variable_value(client, sales_headers, seeded):
    """The insulation-radio path (variable_value-only body) stays require_user."""
    r = client.put(f"/api/bom/{seeded['bom_id']}", headers=sales_headers,
                   json={"variable_value": 0.05})
    assert r.status_code == 200, r.text
    r = client.put(f"/api/bom/{seeded['bom_id']}", headers=sales_headers,
                   json={"variable_value": None})
    assert r.status_code == 200


# ── R2/R4: multi-ratio math + format parity ───────────────────────────────────
def _preview_payload(**over):
    payload = {"result": RESULT_FIXTURE, "dims": DIMS_FIXTURE,
               "trailer_name": FIXED_TT_NAME}
    payload.update(over)
    return payload


def test_excel_multi_ratio_totals_only(client, admin_headers):
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_preview_payload(format="excel", detail="totals",
                                          ratios=[0.35, 0.5, 0.65]))
    assert r.status_code == 200, r.text
    ws, col_a, col_i, flat = _sheet_cells(r.content)
    # zero item rows on the totals-only variant
    assert "FMT FLOOR SHEET" not in flat and "FMT SIDE PANEL" not in flat
    assert "Materials Cost" in col_a and "Margin (10%)" in flat
    # one TOTAL COST @ line per ratio, each = canonical formula output
    for rv, lbl in ((0.35, "35%"), (0.5, "50%"), (0.65, "65%")):
        assert f"TOTAL COST @ {lbl}" in col_a
        assert _expected_total(rv) in col_i
    # a single combined total of two ratios must NOT exist
    assert flat.count("TOTAL COST @") == 3
    assert "Cost per m²" not in flat


def test_excel_single_ratio_replicates_page(client, admin_headers):
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_preview_payload(format="excel", detail="totals",
                                          ratios=[0.55]))
    assert r.status_code == 200, r.text
    _, col_a, col_i, flat = _sheet_cells(r.content)
    assert "Ratio (55%)" in col_a and "TOTAL COST" in col_a
    assert "TOTAL COST @" not in flat            # single ratio → page-shape rows
    assert _expected_total(0.55) in col_i


def test_word_preview_two_ratios(client, admin_headers):
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_preview_payload(format="word", detail="totals",
                                          ratios=[0.35, 0.55]))
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert r.headers["content-disposition"].endswith('.docx"')
    text = _docx_text(r.content)
    assert f"Testing — {FIXED_TT_NAME} (13.6 m)" in text
    assert "Client:  — no client selected —" in text
    assert "TOTAL COST @ 35%" in text and "TOTAL COST @ 55%" in text
    assert "FMT FLOOR SHEET" not in text         # totals-only
    assert "Cost per m²" not in text


def test_word_preview_with_line_items_and_customer(client, admin_headers):
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_preview_payload(format="word", detail="items",
                                          ratios=[0.55],
                                          customer_name="ACME COLD CHAIN"))
    assert r.status_code == 200, r.text
    text = _docx_text(r.content)
    assert "Client:  ACME COLD CHAIN" in text
    assert "FMT FLOOR SHEET" in text and "FMT SIDE PANEL" in text
    # items come BELOW the summary (R2.6)
    assert text.index("TOTAL COST") < text.index("FMT FLOOR SHEET")


def test_pdf_preview_generates(client, admin_headers):
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_preview_payload(format="pdf", detail="totals",
                                          ratios=[0.35, 0.65]))
    assert r.status_code == 200, r.text
    assert r.content[:5] == b"%PDF-"
    assert r.headers["content-disposition"].endswith('.pdf"')


# ── PDF page-1 cover (Michael 8 Aug: the summary spilled onto page 2) ──────────
def _pdf_pages(pdf_bytes: bytes):
    from pypdf import PdfReader
    return [(p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages]


def _wide_payload(n_cats: int, ratios, detail="items"):
    """A body with n_cats DISTINCT categories (category_totals is a dict — repeat
    a name and the count silently collapses) and matching line items."""
    cats = [f"CATEGORY {i:02d}" for i in range(n_cats)]
    items = [{"category": c, "material": f"MAT {c}", "material_code": "",
              "formula": "1", "quantity": 1.0, "unit": "ea", "unit_price": 100.0,
              "waste_pct": 0, "line_cost": 100.0, "last_updated": None}
             for c in cats for _ in range(3)]
    result = dict(RESULT_FIXTURE)
    result["items"] = items
    result["category_totals"] = {c: 300.0 for c in cats}
    return _preview_payload(format="pdf", detail=detail, ratios=list(ratios),
                            result=result)


def test_pdf_page_one_carries_the_whole_cover(client, admin_headers):
    """Heading, client, category totals AND every TOTAL COST line land on page
    1; the line items start on page 2 (they used to share page 1 on a short
    body and push the summary off it on a long one)."""
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_wide_payload(11, (0.35, 0.45, 0.55)))
    assert r.status_code == 200, r.text
    pages = _pdf_pages(r.content)
    assert len(pages) >= 2
    p1 = pages[0]
    assert FIXED_TT_NAME in p1 and "Client:" in p1
    assert "Subtotal (R)" in p1                     # category totals block
    assert "Materials Cost" in p1
    for lbl in ("TOTAL COST @ 35%", "TOTAL COST @ 45%", "TOTAL COST @ 55%"):
        assert lbl in p1.replace("\n", " ") or lbl.replace(" ", "") in p1.replace("\n", "").replace(" ", "")
    assert "Unit Price (R)" not in p1               # items never share the cover
    assert "Unit Price (R)" in pages[1]             # …they start on page 2


def test_pdf_cover_holds_for_a_long_category_list(client, admin_headers):
    """Well past the single-column ceiling the categories lay out two-up rather
    than pushing the summary to page 2."""
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_wide_payload(24, (0.35, 0.55)))
    assert r.status_code == 200, r.text
    p1 = _pdf_pages(r.content)[0]
    assert "Subtotal (R)" in p1 and "Materials Cost" in p1
    assert "CATEGORY 00" in p1 and "CATEGORY 23" in p1   # every category on page 1
    flat = p1.replace("\n", " ").replace(" ", "")
    assert "TOTALCOST@35%" in flat and "TOTALCOST@55%" in flat


def test_pdf_totals_only_is_a_single_page(client, admin_headers):
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_wide_payload(11, (0.55,), detail="totals"))
    assert r.status_code == 200, r.text
    pages = _pdf_pages(r.content)
    assert len(pages) == 1, f"totals-only should be one page, got {len(pages)}"
    assert "Materials Cost" in pages[0] and "MAT CATEGORY 00" not in pages[0]


def test_preview_bad_format_rejected(client, admin_headers):
    r = client.post("/api/export/preview", headers=admin_headers,
                    json=_preview_payload(format="csv"))
    assert r.status_code == 400


def test_preview_writes_nothing(client, admin_headers):
    from app.database import SessionLocal, CalculationRecord
    with SessionLocal() as db:
        before = db.query(CalculationRecord).count()
    for fmt in ("excel", "word", "pdf"):
        r = client.post("/api/export/preview", headers=admin_headers,
                        json=_preview_payload(format=fmt, detail="totals"))
        assert r.status_code == 200
    with SessionLocal() as db:
        assert db.query(CalculationRecord).count() == before


def test_alias_route_pins_excel(client, admin_headers):
    """Old clients POSTing /api/export/excel-preview always get xlsx back,
    even if a format key sneaks into the body."""
    r = client.post("/api/export/excel-preview", headers=admin_headers,
                    json=_preview_payload(format="word"))
    assert r.status_code == 200
    assert r.headers["content-disposition"].endswith('.xlsx"')


# ── approved-side Word + permission gates ─────────────────────────────────────
def test_approved_word_export(client, admin_headers, seeded):
    r = client.get(f"/results/{seeded['rec_id']}/export/word?detail=totals&ratios=0.35,0.65",
                   headers=admin_headers)
    assert r.status_code == 200, r.text
    text = _docx_text(r.content)
    assert f"#{seeded['rec_id']} — {FIXED_TT_NAME} (13.6 m)" in text
    assert "Testing" not in text                 # approved: no Testing heading
    assert "TOTAL COST @ 35%" in text and "TOTAL COST @ 65%" in text


def test_approved_word_gated(client, planner_headers, seeded):
    r = client.get(f"/results/{seeded['rec_id']}/export/word", headers=planner_headers)
    assert r.status_code == 403
    r = client.get(f"/results/{seeded['rec_id']}/export/word")
    assert r.status_code == 401


def test_approved_pdf_with_params(client, admin_headers, seeded):
    r = client.get(f"/results/{seeded['rec_id']}/export/pdf?detail=totals&ratios=0.5",
                   headers=admin_headers)
    assert r.status_code == 200, r.text
    assert r.content[:5] == b"%PDF-"


# ── R6: list payload carries body_length ──────────────────────────────────────
def test_calculations_list_body_length(client, admin_headers, seeded):
    r = client.get("/api/calculations?limit=1000", headers=admin_headers)
    assert r.status_code == 200
    row = next((x for x in r.json() if x["id"] == seeded["rec_id"]), None)
    assert row, "seeded record missing from list"
    assert row["body_length"] == 13.6
