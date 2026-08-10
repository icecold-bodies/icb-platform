"""v1.45 — emailing a costing document from the Preview / Export dialog.

Covers both endpoints (live preview + approved record):
  * the rendered document is ATTACHED, with the right filename/extension per format
    and byte-identical to what the download endpoint would have produced;
  * one recipient, the sender Cc'd from their account email (skipped when blank —
    many accounts still carry email='' since the column only landed in 0030);
  * PREVIEW mail is framed as an internal draft (Michael, 10 Aug) and approved mail
    is not; the optional note rides in the body; totals appear in the body;
  * recipient validation (blank / malformed) and the permission gate;
  * unconfigured SMTP fails LOUDLY (503) rather than silently no-opping — the user
    is waiting on the dialog;
  * emailing a preview still writes NOTHING to the DB.

SMTP is never really opened: `_deliver` is monkeypatched and the EmailMessage is
captured, so these assert on the actual MIME the app built.
"""
import io
import json
import uuid
from datetime import datetime

import pytest

FIXED_RECORD_ID = 900146
FIXED_TT_NAME = "V145 EMAIL BODY"
FIXED_CREATED_AT = datetime(2026, 3, 3, 8, 0, 0)

RESULT_FIXTURE = {
    "items": [
        {"category": "FLOOR", "material": "MAIL FLOOR SHEET", "material_code": "MF-1",
         "formula": "L*W", "quantity": 4.0, "unit": "m2", "unit_price": 100.0,
         "waste_pct": 0, "line_cost": 400.0, "last_updated": None},
    ],
    "category_totals": {"FLOOR": 400.0},
    "grand_total": 400.0,
    "profit_margin": 10,
    "ratio_value": 0.5,
}
DIMS_FIXTURE = {"length": 9.0, "width": 2.5, "height": 2.6}


# ── fixtures ──────────────────────────────────────────────────────────────────
@pytest.fixture(scope="module")
def app_mod():
    import app.main as m
    from starlette.testclient import TestClient
    with TestClient(m.app) as _c:
        yield m


@pytest.fixture(scope="module")
def client(app_mod):
    from starlette.testclient import TestClient
    with TestClient(app_mod.app) as c:
        yield c


def _session_for(username: str) -> dict:
    from app.database import SessionLocal, User, UserSession
    sid = f"v145-{uuid.uuid4().hex[:12]}"
    csrf = f"csrf-{sid}"
    with SessionLocal() as db:
        u = db.query(User).filter_by(username=username).first()
        assert u, f"user {username!r} missing"
        db.merge(UserSession(id=sid, user_id=u.id, role=u.role,
                             expires_at=None, csrf_token=csrf))
        db.commit()
    return {"Cookie": f"session_id={sid}", "X-CSRF-Token": csrf}


@pytest.fixture(scope="module")
def admin_headers(app_mod):
    """Admin, WITH an account email so the sender-Cc path is exercised."""
    from app.database import SessionLocal, User
    with SessionLocal() as db:
        u = db.query(User).filter_by(username="admin").first()
        prior = u.email
        u.email = "sender@icecoldgrp.co.za"
        db.commit()
    headers = _session_for("admin")
    yield headers
    with SessionLocal() as db:
        u = db.query(User).filter_by(username="admin").first()
        u.email = prior
        db.commit()


@pytest.fixture(scope="module")
def full_no_email_headers(app_mod):
    """Role 'full' with a BLANK account email — Nadie's real shape on dev today."""
    from app.database import SessionLocal, User, UserSession
    uname = f"t_full_{uuid.uuid4().hex[:6]}"
    with SessionLocal() as db:
        db.add(User(username=uname, password_hash="x", role="full", email=""))
        db.commit()
    headers = _session_for(uname)
    yield headers
    with SessionLocal() as db:
        db.query(UserSession).filter_by(
            id=headers["Cookie"].split("session_id=")[1]).delete()
        db.query(User).filter_by(username=uname).delete()
        db.commit()


@pytest.fixture(scope="module")
def planner_headers(app_mod):
    from app.database import SessionLocal, User, UserSession
    uname = f"t_planner_{uuid.uuid4().hex[:6]}"
    with SessionLocal() as db:
        db.add(User(username=uname, password_hash="x", role="planner"))
        db.commit()
    headers = _session_for(uname)
    yield headers
    with SessionLocal() as db:
        db.query(UserSession).filter_by(
            id=headers["Cookie"].split("session_id=")[1]).delete()
        db.query(User).filter_by(username=uname).delete()
        db.commit()


@pytest.fixture(scope="module")
def seeded(app_mod):
    from app.database import (SessionLocal, TrailerType, CalculationRecord, User,
                              Customer)
    with SessionLocal() as db:
        admin = db.query(User).filter_by(username="admin").first()
        tt = db.query(TrailerType).filter_by(name=FIXED_TT_NAME).first()
        if not tt:
            tt = TrailerType(name=FIXED_TT_NAME, description="v1.45 email fixture")
            db.add(tt)
            db.flush()
        cust = db.query(Customer).filter_by(name="V145 EMAIL CUSTOMER").first()
        if not cust:
            cust = Customer(name="V145 EMAIL CUSTOMER")
            db.add(cust)
            db.flush()
        rec = db.query(CalculationRecord).filter_by(id=FIXED_RECORD_ID).first()
        if not rec:
            rec = CalculationRecord(
                id=FIXED_RECORD_ID, trailer_type_id=tt.id, user_id=admin.id,
                customer_id=cust.id, quote_number="A99145/03/2026",
                contact_email="attention@customer.co.za",
                dimensions_json=json.dumps(DIMS_FIXTURE),
                result_json=json.dumps(RESULT_FIXTURE),
                created_at=FIXED_CREATED_AT, status="pending", is_repair=False)
            db.add(rec)
        db.commit()
        ids = {"tt_id": tt.id, "rec_id": FIXED_RECORD_ID, "cust_id": cust.id}
    yield ids
    with SessionLocal() as db:
        db.query(CalculationRecord).filter_by(id=FIXED_RECORD_ID).delete()
        db.query(TrailerType).filter_by(id=ids["tt_id"]).delete()
        db.query(Customer).filter_by(id=ids["cust_id"]).delete()
        db.commit()


@pytest.fixture()
def outbox(monkeypatch):
    """Capture the EmailMessage instead of opening SMTP. Also forces a non-empty
    SMTP_URL so the 'not configured' guard doesn't short-circuit the send."""
    import app.services.notifications as notif
    sent: list = []
    monkeypatch.setattr(notif.settings, "SMTP_URL", "smtp://user:pw@mail.test:587",
                        raising=False)
    monkeypatch.setattr(notif.settings, "EMAIL_FROM", "mes@icecoldgrp.co.za",
                        raising=False)
    monkeypatch.setattr(notif, "_deliver", lambda msg, url: sent.append(msg))
    return sent


@pytest.fixture()
def frozen_clock(monkeypatch):
    """Pin the renderers' clock. Every document carries a "Generated {dd Mon YYYY
    HH:MM}" footer, so two renders only match byte-for-byte inside the same
    minute — comparing them unpinned is a coin-flip that passes alone and fails
    in a full run (which is exactly how it first showed up)."""
    import app.routers.exports as ex
    fixed = datetime(2026, 3, 3, 9, 30, 0)

    class _FixedDT(datetime):        # subclass: fromisoformat/timedelta still work
        @classmethod
        def now(cls, tz=None):
            return fixed if tz is None else fixed.replace(tzinfo=tz)

    monkeypatch.setattr(ex, "datetime", _FixedDT)
    return fixed


def _attachment(msg):
    for part in msg.iter_attachments():
        return part
    return None


def _body_text(msg) -> str:
    for part in msg.walk():
        if part.get_content_type() == "text/plain" and not part.get_filename():
            return part.get_content()
    return ""


def _preview_payload(**over):
    payload = {"result": RESULT_FIXTURE, "dims": DIMS_FIXTURE,
               "trailer_name": FIXED_TT_NAME, "to": "nadie@icecoldgrp.co.za"}
    payload.update(over)
    return payload


# ── preview email ─────────────────────────────────────────────────────────────
def test_preview_email_attaches_the_document(client, admin_headers, outbox):
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf", detail="totals", ratios=[0.5]))
    assert r.status_code == 200, r.text
    assert r.json()["to"] == "nadie@icecoldgrp.co.za"
    assert len(outbox) == 1
    msg = outbox[0]
    assert msg["To"] == "nadie@icecoldgrp.co.za"
    assert msg["Cc"] == "sender@icecoldgrp.co.za"          # sender copied in
    att = _attachment(msg)
    assert att is not None, "no attachment on the message"
    assert att.get_filename().endswith(".pdf")
    assert att.get_payload(decode=True)[:5] == b"%PDF-"


def test_preview_email_is_framed_as_a_draft(client, admin_headers, outbox):
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="excel", note="Checked the floor twice."))
    assert r.status_code == 200, r.text
    msg = outbox[0]
    assert "DRAFT" in msg["Subject"] and "not a quotation" in msg["Subject"].lower()
    body = _body_text(msg)
    assert "INTERNAL DRAFT" in body
    assert "no quote number" in body
    assert "Checked the floor twice." in body               # the optional note rides along
    assert "TOTAL COST" in body                             # totals summarised in the body
    assert _attachment(msg).get_filename().endswith(".xlsx")


def _ooxml_digest(blob: bytes) -> str:
    """Content digest of an OOXML (.docx/.xlsx) package: sha256 over each zip
    member's NAME + DECOMPRESSED bytes, in sorted order, skipping docProps/core.xml.

    Raw bytes can't be compared directly. python-docx/openpyxl hand each part to
    zipfile.writestr(name, blob), which stamps the entry with the wall clock at
    save time (2-second granularity) — so two renders of identical content differ
    in zip metadata whenever they straddle a bucket. That is what made a naive
    equality assertion pass locally and fail on the Windows CI leg. core.xml is
    excluded for the same reason (it carries created/modified). This is the same
    normalization the v1.44 saved-export lock used.
    """
    import hashlib
    import zipfile
    zf = zipfile.ZipFile(io.BytesIO(blob))
    h = hashlib.sha256()
    for name in sorted(zf.namelist()):
        if name == "docProps/core.xml":
            continue
        h.update(name.encode())
        h.update(b"\x00")
        h.update(zf.read(name))
    return h.hexdigest()


def test_preview_email_matches_the_download(client, admin_headers, outbox,
                                            frozen_clock):
    """The emailed attachment IS the document the user would have downloaded —
    same renderer, same options, same content. The clock is pinned so the
    "Generated {time}" footer inside the document matches too; the digest then
    ignores only the zip's own save-time metadata."""
    opts = {"format": "word", "detail": "items", "ratios": [0.35, 0.55]}
    dl = client.post("/api/export/preview", headers=admin_headers,
                     json=_preview_payload(**opts))
    assert dl.status_code == 200, dl.text
    em = client.post("/api/export/preview/email", headers=admin_headers,
                     json=_preview_payload(**opts))
    assert em.status_code == 200, em.text
    attached = _attachment(outbox[0]).get_payload(decode=True)
    assert _ooxml_digest(attached) == _ooxml_digest(dl.content)
    # …and the text really is the document, not an empty shell.
    assert f"Testing — {FIXED_TT_NAME}" in _docx_text(attached)


def _docx_text(blob: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "|".join(parts)


def test_preview_email_writes_nothing(client, admin_headers, outbox):
    from app.database import SessionLocal, CalculationRecord
    with SessionLocal() as db:
        before = db.query(CalculationRecord).count()
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf"))
    assert r.status_code == 200
    with SessionLocal() as db:
        assert db.query(CalculationRecord).count() == before


def test_sender_cc_skipped_when_account_has_no_email(client, full_no_email_headers, outbox):
    r = client.post("/api/export/preview/email", headers=full_no_email_headers,
                    json=_preview_payload(format="excel"))
    assert r.status_code == 200, r.text
    assert r.json()["cc"] is None
    assert outbox[0]["Cc"] is None                          # sent, just without a copy


# ── approved email ────────────────────────────────────────────────────────────
def test_approved_email_is_not_a_draft(client, admin_headers, seeded, outbox):
    r = client.post(f"/results/{seeded['rec_id']}/export/email", headers=admin_headers,
                    json={"to": "nadie@icecoldgrp.co.za", "format": "pdf",
                          "detail": "totals", "ratios": [0.5]})
    assert r.status_code == 200, r.text
    msg = outbox[0]
    assert "DRAFT" not in msg["Subject"]
    assert "A99145/03/2026" in msg["Subject"]               # quote number in the subject
    body = _body_text(msg)
    assert "INTERNAL DRAFT" not in body
    assert "V145 EMAIL CUSTOMER" in body                    # client named in the body
    assert _attachment(msg).get_filename().endswith(".pdf")


def test_approved_email_defaults_to_pdf(client, admin_headers, seeded, outbox):
    r = client.post(f"/results/{seeded['rec_id']}/export/email", headers=admin_headers,
                    json={"to": "nadie@icecoldgrp.co.za"})
    assert r.status_code == 200, r.text
    assert _attachment(outbox[0]).get_filename().endswith(".pdf")


def test_approved_email_404_for_unknown_record(client, admin_headers, outbox):
    r = client.post("/results/99999999/export/email", headers=admin_headers,
                    json={"to": "nadie@icecoldgrp.co.za"})
    assert r.status_code == 404


# ── guards ────────────────────────────────────────────────────────────────────
@pytest.mark.parametrize("bad", ["", "   ", "not-an-email", "a@b", "two@addr.com, x@y.com"])
def test_bad_recipient_rejected(client, admin_headers, outbox, bad):
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf", to=bad))
    assert r.status_code == 400, r.text
    assert not outbox, "nothing may be sent when the address is rejected"


# ── v1.45.1 internal-only allowlist ───────────────────────────────────────────
# Michael 10 Aug, after a test send put a customer's priced costing on an outside
# domain: a costing may only be emailed to an internal mailbox.

@pytest.mark.parametrize("addr", [
    "nadie@icecoldgrp.co.za",
    "NADIE@IceColdGrp.CO.ZA",            # case-insensitive
    "  burt@icecoldgrp.co.za  ",         # trimmed
    "micger123@gmail.com",               # the BA/owner, listed verbatim
])
def test_internal_recipients_allowed(client, admin_headers, outbox, addr):
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf", to=addr))
    assert r.status_code == 200, r.text
    assert len(outbox) == 1
    assert outbox[0]["To"].strip().lower() == addr.strip().lower()


@pytest.mark.parametrize("addr", [
    "x@y.co.za",                              # the address that caused the incident
    "buyer@customer.co.za",
    "mandy@360degreescarriers.co.za",         # a real customer contact
    "someone@gmail.com",                      # not the listed owner address
    "finance@noticecoldgrp.co.za",            # endswith() would have waved this through
    "x@icecoldgrp.co.za.attacker.com",        # …and this
    "x@sub.icecoldgrp.co.za",                 # sub-domains are NOT inherited
])
def test_external_recipients_refused(client, admin_headers, outbox, addr):
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf", to=addr))
    assert r.status_code == 403, f"{addr} was not refused: {r.status_code} {r.text[:200]}"
    assert "internal" in r.json()["detail"].lower()
    assert not outbox, f"{addr} must not receive anything"


def test_approved_email_also_refuses_external(client, admin_headers, seeded, outbox):
    """The approved path pre-fills from the customer's Attention contact, so it is
    the likeliest place to send a priced costing outside — same gate."""
    r = client.post(f"/results/{seeded['rec_id']}/export/email", headers=admin_headers,
                    json={"to": "attention@customer.co.za", "format": "pdf"})
    assert r.status_code == 403
    assert not outbox


def test_sender_cc_dropped_when_the_account_is_external(client, outbox, app_mod):
    """A user whose account carries an outside address must not become a
    side-channel: the send proceeds, the Cc is dropped."""
    from app.database import SessionLocal, User, UserSession
    uname = f"t_ext_{uuid.uuid4().hex[:6]}"
    with SessionLocal() as db:
        db.add(User(username=uname, password_hash="x", role="full",
                    email="someone@outside.example"))
        db.commit()
    headers = _session_for(uname)
    try:
        r = client.post("/api/export/preview/email", headers=headers,
                        json=_preview_payload(format="pdf", to="nadie@icecoldgrp.co.za"))
        assert r.status_code == 200, r.text
        assert r.json()["cc"] is None
        assert outbox[0]["Cc"] is None
    finally:
        with SessionLocal() as db:
            db.query(UserSession).filter_by(
                id=headers["Cookie"].split("session_id=")[1]).delete()
            db.query(User).filter_by(username=uname).delete()
            db.commit()


def test_empty_allowlist_fails_closed(client, admin_headers, outbox, monkeypatch):
    """Blanking the config must shut the gate, not open it — the classic
    misconfiguration that turns a deny-list into an allow-all."""
    from app.config import settings
    monkeypatch.setattr(settings, "COSTING_EMAIL_ALLOWED_DOMAINS", "", raising=False)
    monkeypatch.setattr(settings, "COSTING_EMAIL_ALLOWED_ADDRESSES", "", raising=False)
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf", to="nadie@icecoldgrp.co.za"))
    assert r.status_code == 403
    assert not outbox


def test_policy_endpoint_reports_the_rule(client, admin_headers):
    r = client.get("/api/export/email-policy", headers=admin_headers)
    assert r.status_code == 200, r.text
    d = r.json()
    assert "icecoldgrp.co.za" in d["domains"]
    assert "micger123@gmail.com" in d["addresses"]
    assert "internal" in d["message"].lower()


def test_policy_endpoint_needs_a_session(client):
    assert client.get("/api/export/email-policy").status_code == 401


def test_email_gated_on_the_format_permission(client, planner_headers, outbox):
    r = client.post("/api/export/preview/email", headers=planner_headers,
                    json=_preview_payload(format="pdf"))
    assert r.status_code == 403
    assert not outbox


def test_email_requires_a_session(client, outbox):
    r = client.post("/api/export/preview/email", json=_preview_payload(format="pdf"))
    assert r.status_code == 401
    assert not outbox


def test_bad_format_rejected(client, admin_headers, outbox):
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="csv"))
    assert r.status_code == 400
    assert not outbox


def test_unconfigured_smtp_fails_loudly(client, admin_headers, monkeypatch):
    """No SMTP → 503 with a plain-English message, NOT a silent success. The
    sibling helpers log-and-continue; this one must not (the user is waiting)."""
    import app.services.notifications as notif
    monkeypatch.setattr(notif.settings, "SMTP_URL", "", raising=False)
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf"))
    assert r.status_code == 503, r.text
    assert "not configured" in r.json()["detail"].lower()


def test_relay_failure_surfaces_as_502(client, admin_headers, monkeypatch):
    import app.services.notifications as notif
    monkeypatch.setattr(notif.settings, "SMTP_URL", "smtp://user:pw@mail.test:587",
                        raising=False)

    def _boom(msg, url):
        raise OSError("mailbox unavailable")
    monkeypatch.setattr(notif, "_deliver", _boom)
    r = client.post("/api/export/preview/email", headers=admin_headers,
                    json=_preview_payload(format="pdf"))
    assert r.status_code == 502, r.text
    assert "could not be sent" in r.json()["detail"].lower()
