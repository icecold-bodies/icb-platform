"""WO v1.47 lane B — the END USER (the customer's customer), end to end (marker J147EU).

Nadie's loop in ONE browser session, through the costings-page calculator embed:

  a customer with an end user on file → the "End user (optional)" picker appears under
  Attention (and does NOT auto-select — an end user is a claim about who the body is
  for) → pick it → Approve & Save → the costing detail page shows the End user row →
  export Excel / Word / PDF and each carries "End user: {company}" + "End user contact:
  {name}" → rename the end-user record → the approved costing STILL prints the original.

Plus the two negatives that matter: a customer with no end users shows the empty state
with a reachable "+ Add now" (whose quick-add comes back selected), and a costing saved
with NO end user produces documents with no end-user text anywhere.

WHY ONE BIG TEST: the loop is stateful — pick, approve, then read the same record back
through three renderers and a rename. Splitting it gives each phase a fresh browser
profile that has to re-derive the last one's state (the banked v1.45 lesson).

Selector policy: Jinja pages by element ID; React by data-testid. No wait_for_function
(the app CSP has no unsafe-eval).
"""
from __future__ import annotations

import io

import pytest
from playwright.sync_api import Page, expect

from _common import role_session, shot  # noqa: E402  (sys.path set in conftest)

T = 15_000
JOURNEY = "end_user"
MARK = "J147EU"
FULL_USER = "journey_full_enduser"     # role 'full' = Nadie

_EMBED = "iframe[title='Calculator (live costing app)']"


def _purge(db) -> None:
    from sqlalchemy import text
    db.execute(text(
        "DELETE FROM icb_costings.calculations WHERE trailer_type_id IN "
        "(SELECT id FROM icb_costings.trailer_types WHERE name LIKE :m)"), {"m": f"{MARK}%"})
    db.execute(text(
        "DELETE FROM icb_costings.calculations cal USING icb_costings.customers c "
        "WHERE cal.customer_id = c.id AND c.name LIKE :m"), {"m": f"{MARK}%"})
    db.execute(text(
        "DELETE FROM icb_costings.customer_end_users eu USING icb_costings.customers c "
        "WHERE eu.customer_id = c.id AND c.name LIKE :m"), {"m": f"{MARK}%"})
    db.execute(text("DELETE FROM icb_costings.customers WHERE name LIKE :m"), {"m": f"{MARK}%"})
    db.execute(text(
        "DELETE FROM icb_costings.bill_of_materials WHERE material_id IN "
        "(SELECT id FROM icb_costings.materials WHERE name LIKE :m)"), {"m": f"{MARK}%"})
    db.execute(text("DELETE FROM icb_costings.materials WHERE name LIKE :m"), {"m": f"{MARK}%"})
    db.execute(text("DELETE FROM icb_costings.trailer_types WHERE name LIKE :m"), {"m": f"{MARK}%"})
    db.commit()


@pytest.fixture(scope="module")
def staged():
    """One calculable marker body, a RESELLER customer with one end user on file, a
    BARE customer with none, and Nadie's role-'full' journey user."""
    from app.database import (BillOfMaterial, Customer, CustomerEndUser, Material,
                              SessionLocal, TrailerType, User)
    created_user = False
    with SessionLocal() as db:
        _purge(db)

        tt = TrailerType(name=f"{MARK} CHILLER", is_active=True,
                         default_length=7.5, default_width=2.5, default_height=2.6)
        db.add(tt)
        db.flush()
        mat = Material(name=f"{MARK} PANEL", unit_of_measure="m2", price_per_unit=1000.0)
        db.add(mat)
        db.flush()
        db.add(BillOfMaterial(trailer_type_id=tt.id, material_id=mat.id,
                              formula_expression="1", waste_percentage=0,
                              bom_section="PANELS", sort_order=1))

        reseller = Customer(name=f"{MARK} Reseller Ltd", bp_code=f"{MARK}1", is_active=True)
        bare = Customer(name=f"{MARK} Bare Ltd", bp_code=f"{MARK}2", is_active=True)
        db.add_all([reseller, bare])
        db.flush()
        acme = CustomerEndUser(customer_id=reseller.id, company_name=f"{MARK} ACME Foods",
                               contact_name="Thabo Nkosi", contact_role="Fleet",
                               contact_email="thabo@acme.co", contact_telephone="011 999",
                               is_primary=True, active=True)
        db.add(acme)
        if db.query(User).filter_by(username=FULL_USER).first() is None:
            db.add(User(username=FULL_USER, password_hash="x", role="full"))
            created_user = True
        db.commit()
        ids = {"trailer": tt.id, "name": tt.name, "reseller": reseller.id,
               "bare": bare.id, "acme": acme.id}

    yield ids

    with SessionLocal() as db:
        _purge(db)
        if created_user:
            u = db.query(User).filter_by(username=FULL_USER).first()
            if u is not None:
                db.delete(u)
        db.commit()


def _select_trailer_and_wait_bom(page: Page, frame, trailer_id: str,
                                 attempts: int = 4, expect_length: str = "7.5") -> None:
    """select_option is one-shot: while the embed is still settling, the
    change→loadBOM→fetch chain can be lost and no BOM row renders (banked embed-journey
    flake class). Re-select, bounded. Waits for ATTACHED, not visible — the post-calc
    table regroups by category and remembers a COLLAPSED state. The dimension check is in
    the same loop because a body's DEFAULT dims land only once /api/trailers has resolved;
    pick before that and the inputs keep the template's own values."""
    row = frame.locator("[data-material-id]").first
    length = frame.locator("#f-length")
    last = ""
    for _ in range(attempts):
        frame.locator("#trailer-select").select_option(trailer_id)
        try:
            expect(row).to_be_attached(timeout=8_000)
            expect(length).to_have_value(expect_length, timeout=6_000)
            return
        except AssertionError:
            last = length.input_value()
            continue
    shot(page, "zz-body-never-settled", journey=JOURNEY)
    raise AssertionError(
        f"body did not settle after {attempts} selects "
        f"(rows attached={row.count()}, length={last!r}, wanted {expect_length!r})")


def _wait_calc_idle(page: Page, frame, quiet_ms: int = 1_200, timeout_s: int = 60) -> None:
    """`expect(#approve-btn).to_be_enabled()` alone is a LEVEL check and is not enough:
    loadBOM and every dimension edit go through scheduleCalc()'s 700 ms debounce, and
    runCalc() disables the button synchronously only once that timer fires — in the gap
    the button is still enabled from the PREVIOUS calc. Wait for it to STAY enabled across
    a window longer than the debounce (the banked level-vs-edge lesson)."""
    import time
    approve = frame.locator("#approve-btn")
    deadline = time.monotonic() + timeout_s
    while time.monotonic() < deadline:
        expect(approve).to_be_enabled(timeout=30_000)
        page.wait_for_timeout(quiet_ms)
        if approve.is_enabled():
            return
    raise AssertionError(
        f"the calculator never went idle within {timeout_s}s "
        "(#approve-btn kept flipping back to disabled)")


def _open_embed(page: Page, live_server: str, trailer_id: str):
    """Nadie's surface: /mes-app/costings/new with the calculator embed."""
    role_session(page, FULL_USER, base=live_server)
    page.goto("/mes-app/costings/new")
    frame = page.frame_locator(_EMBED)
    expect(frame.locator("#trailer-select")).to_be_visible(timeout=30_000)
    _select_trailer_and_wait_bom(page, frame, trailer_id)
    _wait_calc_idle(page, frame)
    return frame


def _approve(page: Page, frame) -> None:
    """Save the costing and wait on the /api/approve RESPONSE, not the success toast.

    Two reasons the toast is the wrong signal here. It is transient, and the successful
    save also postMessages `mes:costing-saved` to the shell, which refetches the dashboard
    under the embed — so the toast can be gone before an assertion reaches it. The response
    is edge-triggered and carries the status, which is what we actually want to know.

    A customer that already has a costing for this body type opens the DUPLICATE modal
    first, and approve does not fire until a branch is chosen — that is why this journey's
    later saves would otherwise hang with no request at all. Take the independent
    "Save as new costing" branch: these assertions are about the end user, not versioning.
    """
    with page.expect_response(
            lambda r: "/api/approve" in r.url and r.request.method == "POST",
            timeout=30_000) as ri:
        frame.locator("#approve-btn").click()
        no_cust = frame.locator("#modal-no-customer:not(.hidden)")
        try:
            expect(no_cust).to_be_visible(timeout=2_000)
            frame.locator("#modal-no-customer .btn-outline").click()
        except AssertionError:
            pass    # a customer was attached — nothing to dismiss
        dup = frame.locator("#modal-duplicate:not(.hidden)")
        try:
            expect(dup).to_be_visible(timeout=4_000)
            frame.locator("#dup-savenew-btn").click()
        except AssertionError:
            pass    # first costing for this customer + body type
    assert ri.value.status == 200, f"/api/approve returned {ri.value.status}"


def _open_detail(page: Page, quote: str, attempts: int = 3):
    """Open CostingDetail — keyed by QUOTE NUMBER, not record id (/costings/:quote); the
    /costings/results/:id route is a different screen.

    The quote number MUST be percent-encoded. Real numbers carry slashes ("J17/08/2026"),
    so a raw path silently fails to match the single-segment :quote route, the SPA falls
    through, and the assertion then reads whatever page it landed on — /mes-app/production,
    in the run that caught this — which looks nothing like "costing not found". The app's
    own links encode it and CostingDetail decodeURIComponent()s it back.

    The bounded retry then covers the genuine timing case: a successful approve
    postMessages `mes:costing-saved` to the shell, which refetches underneath the embed.
    The failure message names the URL actually reached, so a route miss can never again be
    mistaken for a missing record.
    """
    from urllib.parse import quote as _urlquote
    card = page.get_by_text("Quotation / Configuration Overview").first
    for _ in range(attempts):
        page.goto(f"/mes-app/costings/{_urlquote(quote, safe='')}")
        try:
            expect(card).to_be_visible(timeout=8_000)
            return
        except AssertionError:
            continue
    shot(page, "zz-detail-never-opened", journey=JOURNEY)
    raise AssertionError(
        f"the costing detail for {quote!r} never rendered — page is at {page.url!r}")


def _latest_record(trailer_id: int):
    from app.database import CalculationRecord, SessionLocal
    with SessionLocal() as db:
        return (db.query(CalculationRecord)
                  .filter_by(trailer_type_id=trailer_id)
                  .order_by(CalculationRecord.id.desc()).first())


def _export_text(page: Page, base: str, record_id: int, fmt: str) -> str:
    """Pull an export through the REAL route with the browser's own session, then read the
    bytes with the matching parser. Asserting on parsed content (not on a download event)
    is what actually proves the line reached the document."""
    resp = page.request.get(f"{base}/results/{record_id}/export/{fmt}")
    assert resp.status == 200, f"{fmt} export returned {resp.status}"
    body = resp.body()
    if fmt == "excel":
        import openpyxl
        ws = openpyxl.load_workbook(io.BytesIO(body)).active
        return "|".join(str(c.value) for row in ws.iter_rows()
                        for c in row if c.value is not None)
    if fmt == "word":
        from docx import Document
        doc = Document(io.BytesIO(body))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                parts += [cell.text for cell in row.cells]
        return "|".join(parts)
    from pypdf import PdfReader
    return " ".join((p.extract_text() or "") for p in
                    PdfReader(io.BytesIO(body)).pages).replace("\n", " ")


# ── 1. Nadie's whole loop, in ONE browser session ────────────────────────────
def test_end_user_picked_saved_and_printed(page: Page, live_server: str, staged) -> None:
    frame = _open_embed(page, live_server, str(staged["trailer"]))

    # No customer yet → the block is hidden entirely (there is nothing it could list).
    expect(frame.locator("#end-user-block")).to_be_hidden()

    # Reseller selected → the picker appears, populated, and NOTHING auto-selects.
    with page.expect_response(
            lambda r: f"/api/customers/{staged['reseller']}/end-users" in r.url, timeout=T):
        frame.locator("#cust-select").select_option(str(staged["reseller"]))
    expect(frame.locator("#end-user-block")).to_be_visible(timeout=T)
    expect(frame.locator("#end-user-select")).to_be_enabled(timeout=T)
    assert (frame.locator("#end-user-select").input_value() or "") == "", \
        "an end user must never auto-select — it is a claim about who the body is for"
    options = frame.locator("#end-user-select option").all_inner_texts()
    assert any("ACME Foods" in o and "Thabo Nkosi" in o for o in options), options
    shot(page, "01-picker-under-attention", journey=JOURNEY)

    # …and it sits BELOW the Attention picker (Default 3 — the ratified order).
    att = frame.locator("#contact-block").bounding_box()
    eu = frame.locator("#end-user-block").bounding_box()
    assert att and eu and eu["y"] > att["y"], "the End user block must sit under Attention"

    # Pick it, wait for the calculator to settle, then save.
    frame.locator("#end-user-select").select_option(str(staged["acme"]))
    _wait_calc_idle(page, frame)
    _approve(page, frame)

    rec = _latest_record(staged["trailer"])
    assert rec is not None, "no costing row was saved"
    assert rec.end_user_id == staged["acme"]
    assert rec.end_user_company == f"{MARK} ACME Foods"
    assert rec.end_user_contact_name == "Thabo Nkosi"
    # The sibling contact snapshot is untouched by this WO.
    assert rec.end_user_contact_email == "thabo@acme.co"

    # ── the detail page shows the End user row ───────────────────────────────
    assert rec.quote_number, "the approved costing has no quote number to open"
    _open_detail(page, rec.quote_number)
    expect(page.get_by_text("End user", exact=True).first).to_be_visible(timeout=T)
    expect(page.get_by_text(f"{MARK} ACME Foods").first).to_be_visible(timeout=T)
    shot(page, "02-detail-end-user-row", journey=JOURNEY)

    # ── all three export formats carry both lines ────────────────────────────
    for fmt in ("excel", "word", "pdf"):
        text = _export_text(page, live_server, rec.id, fmt)
        assert f"End user: {MARK} ACME Foods" in text, f"{fmt}: end-user company line missing"
        assert "End user contact: Thabo Nkosi" in text, f"{fmt}: end-user contact line missing"

    # ── Default 6: reopening the costing restores the selected end user ──────
    # Straight at the Jinja page (no embed) — the same template the shell iframes.
    # Restore rides setCustomer(customerId, contactId, endUserId), the one path the
    # contact picker already used, so edit and duplicate get it together.
    page.goto(f"/mes/calculator?edit={rec.id}")
    expect(page.locator("#end-user-select")).to_be_enabled(timeout=30_000)
    expect(page.locator("#end-user-select")).to_have_value(str(staged["acme"]), timeout=T)
    shot(page, "07-edit-reopen-restores-end-user", journey=JOURNEY)

    # ── rename the record → the approved costing still prints the original ───
    from app.database import CustomerEndUser, SessionLocal
    with SessionLocal() as db:
        acme = db.get(CustomerEndUser, staged["acme"])
        acme.company_name = f"{MARK} ACME RENAMED"
        acme.contact_name = "Someone Else"
        db.commit()

    for fmt in ("excel", "word", "pdf"):
        text = _export_text(page, live_server, rec.id, fmt)
        assert f"End user: {MARK} ACME Foods" in text, \
            f"{fmt}: the snapshot was not frozen — a later rename rewrote quote history"
        assert "RENAMED" not in text and "Someone Else" not in text, f"{fmt}: live values leaked in"
    shot(page, "03-snapshot-frozen", journey=JOURNEY)


# ── 2. The empty state, and the inline quick-add ─────────────────────────────
def test_customer_without_end_users_can_add_one_inline(page: Page, live_server: str,
                                                       staged) -> None:
    frame = _open_embed(page, live_server, str(staged["trailer"]))

    with page.expect_response(
            lambda r: f"/api/customers/{staged['bare']}/end-users" in r.url, timeout=T):
        frame.locator("#cust-select").select_option(str(staged["bare"]))
    # Visible (so "+ Add now" is reachable) but with nothing to pick.
    expect(frame.locator("#end-user-block")).to_be_visible(timeout=T)
    expect(frame.locator("#end-user-empty")).to_be_visible(timeout=T)
    expect(frame.locator("#end-user-select")).to_be_disabled()
    shot(page, "04-empty-state-add-now", journey=JOURNEY)

    frame.locator("#end-user-empty a").click()
    expect(frame.locator("#end-user-add-form")).to_be_visible(timeout=T)
    frame.locator("#end-user-new-company").fill(f"{MARK} Fresh Foods")
    frame.locator("#end-user-new-name").fill("Lerato M")
    with page.expect_response(
            lambda r: r.url.endswith(f"/api/customers/{staged['bare']}/end-users")
            and r.request.method == "POST", timeout=T) as ri:
        frame.locator("#end-user-add-form button:has-text('Save end user')").click()
    assert ri.value.status == 200, f"quick-add returned {ri.value.status}"

    # The one just added IS the intended pick, so it comes back selected.
    expect(frame.locator("#end-user-select")).to_be_enabled(timeout=T)
    picked = frame.locator("#end-user-select option:checked").text_content() or ""
    assert "Fresh Foods" in picked, f"quick-added end user not selected: {picked!r}"
    shot(page, "05-quickadd-selected", journey=JOURNEY)

    # Switching customers re-populates from the other one (no stale carry-over).
    with page.expect_response(
            lambda r: f"/api/customers/{staged['reseller']}/end-users" in r.url, timeout=T):
        frame.locator("#cust-select").select_option(str(staged["reseller"]))
    opts = frame.locator("#end-user-select option").all_inner_texts()
    assert not any("Fresh Foods" in o for o in opts), f"stale carry-over: {opts}"


# ── 3. The optional path: no end user → documents unchanged ──────────────────
def test_costing_without_an_end_user_prints_no_end_user_text(page: Page, live_server: str,
                                                             staged) -> None:
    frame = _open_embed(page, live_server, str(staged["trailer"]))
    with page.expect_response(
            lambda r: f"/api/customers/{staged['reseller']}/end-users" in r.url, timeout=T):
        frame.locator("#cust-select").select_option(str(staged["reseller"]))
    expect(frame.locator("#end-user-select")).to_be_enabled(timeout=T)
    # Deliberately leave it on "— No end user —".
    _wait_calc_idle(page, frame)
    _approve(page, frame)

    rec = _latest_record(staged["trailer"])
    assert rec is not None and rec.end_user_id is None, "an end user was saved anyway"
    assert rec.end_user_company is None

    for fmt in ("excel", "word", "pdf"):
        text = _export_text(page, live_server, rec.id, fmt)
        assert "End user" not in text, \
            f"{fmt}: a costing with no end user must print no end-user text at all"
        assert f"{MARK} Reseller Ltd" in text, f"{fmt}: the client line went missing"


# ── 4. The admin surface: end users are managed beside contacts ──────────────
def test_end_users_are_managed_on_the_customers_admin_page(page: Page, live_server: str,
                                                           staged) -> None:
    """Default 4 — the SAME screen as customer contacts, a second section, no new page."""
    role_session(page, "admin", base=live_server)
    page.goto("/mes-app/admin/customers")
    page.get_by_test_id("customers-search").fill(f"{MARK} Reseller")
    row = page.get_by_test_id("customer-row").first
    expect(row).to_be_visible(timeout=T)
    row.click()

    contacts = page.get_by_test_id("contacts-panel")
    end_users = page.get_by_test_id("end-users-panel")
    expect(contacts).to_be_visible(timeout=T)
    expect(end_users).to_be_visible(timeout=T)
    # Second section, under contacts — not a replacement for them.
    cb, eb = contacts.bounding_box(), end_users.bounding_box()
    assert cb and eb and eb["y"] > cb["y"], "End users must sit under Contacts"
    # Match on the marker prefix, not the company name: the snapshot test above renames
    # this very record on purpose, and these tests share one module-scoped fixture.
    expect(end_users.get_by_test_id("end-user-row").first).to_contain_text(MARK)
    shot(page, "06-customers-admin-end-users-section", journey=JOURNEY)

    # Add one from here too (the admin path, not the calculator quick-add).
    page.get_by_test_id("end-user-add").click()
    page.get_by_test_id("end-user-add-company").fill(f"{MARK} Admin Added Co")
    page.get_by_test_id("end-user-add-name").fill("Admin Person")
    page.get_by_test_id("end-user-add-save").click()
    expect(end_users.get_by_test_id("end-user-row")
           .filter(has_text=f"{MARK} Admin Added Co")).to_be_visible(timeout=T)
